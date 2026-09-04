import os
import sys
import tempfile
import copy
import re
import time
import base64
import requests
from flask import Flask, request, jsonify, send_file, send_from_directory
from flask_cors import CORS
from PIL import Image
from google import genai
from google.genai import types
from pydantic import BaseModel, Field
from typing import List, Optional
import docx
from docx.shared import Pt
import openpyxl
from openpyxl.drawing.spreadsheet_drawing import OneCellAnchor, AnchorMarker
from openpyxl.drawing.xdr import XDRPositiveSize2D
from openpyxl.utils.units import pixels_to_EMU
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from bahttext import bahttext
from dotenv import load_dotenv

load_dotenv()

app = Flask(__name__, static_folder='static', static_url_path='')
CORS(app, resources={r"/*": {"origins": "*"}}, supports_credentials=True)

# ==========================================
# Serve Single Page Application (SPA) Frontend
# ==========================================
@app.route('/')
def serve_index():
    if os.path.exists(os.path.join(app.static_folder, 'index.html')):
        return send_from_directory(app.static_folder, 'index.html')
    return "Backend running. Please build frontend first.", 200

@app.errorhandler(404)
def spa_fallback(e):
    if request.path.startswith('/api/'):
        return jsonify({"error": "Not Found"}), 404
    if os.path.exists(os.path.join(app.static_folder, 'index.html')):
        return send_from_directory(app.static_folder, 'index.html')
    return "Frontend build not found.", 404

# API Key Placeholder
HARDCODED_API_KEY = os.environ.get("API_KEY", "")
# ==========================================
# Contacts Data from Excel (All Sheets)
# ==========================================
_contacts_cache = None

def load_contacts_from_excel():
    """Parse ALL sheets from the organization Excel file.
    Each sheet has the same structure:
    - Row 1: Office/Department title (merged, col A only)
    - Row 2: Headers (ชื่อ-นามสกุล, ตำแหน่ง, หมายเลขโทรศัพท์, ..., Email)
    - Row 3: Sub-headers (มือถือ, ตั้งโต๊ะ)
    - Row 4+: Data rows
      - Section headers: col A has value, col B is None
      - Staff rows: col A has name, col B has position
    """
    global _contacts_cache
    xlsx_path = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                             'data', 'รายชื่อติดต่อบุคลากรภายใน สทอภ. ตามโครงสร.xlsx')
    if not os.path.exists(xlsx_path):
        print(f"[Contacts] Excel file not found: {xlsx_path}")
        _contacts_cache = []
        return []
    
    try:
        wb = openpyxl.load_workbook(xlsx_path, read_only=True, data_only=True)
    except Exception as e:
        print(f"[Contacts] Failed to load Excel: {e}")
        _contacts_cache = []
        return []
    
    all_contacts = []
    
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        
        # Get office name from row 1 col A
        office_name = None
        row1_val = ws.cell(row=1, column=1).value
        if row1_val:
            # Clean up newlines and extract short name
            office_name = str(row1_val).split('\n')[0].strip()
        
        current_section = office_name or sheet_name
        current_section_head = None  # Name of the head of current section
        section_staff = []  # Buffer to assign section_head retroactively
        
        for row in ws.iter_rows(min_row=2, max_row=ws.max_row, max_col=6):
            vals = [c.value for c in row]
            
            # Skip fully empty rows
            if all(v is None for v in vals):
                continue
            
            col_a = str(vals[0]).strip() if vals[0] else None
            col_b = str(vals[1]).strip() if vals[1] else None
            
            if not col_a:
                continue
            
            # Skip header rows
            if col_a == 'ชื่อ-นามสกุล':
                continue
            
            # Section header: col A has value, col B is None or empty
            if not col_b:
                # Finalize previous section: assign section_head
                for s in section_staff:
                    s['section_head'] = current_section_head
                all_contacts.extend(section_staff)
                
                # Start new section
                current_section = col_a
                current_section_head = None
                section_staff = []
                continue
            
            # Staff row
            name = col_a
            position = col_b
            mobile = str(vals[2]).strip() if vals[2] else None
            desk = str(vals[3]).strip() if vals[3] else None
            email = str(vals[4]).strip() if vals[4] else None
            
            is_head = any(kw in position for kw in [
                'หัวหน้าฝ่าย', 'หัวหน้างาน', 'ผู้อำนวยการ',
                'ผช.ผอ', 'รก.ผอ', 'รก.หน'
            ])
            
            # Extract nickname from parentheses in name
            nickname = None
            nick_match = re.search(r'\(([^)]+)\)\s*$', name)
            if nick_match:
                nickname = nick_match.group(1)
            
            contact = {
                'name': name,
                'nickname': nickname,
                'position': position,
                'mobile': mobile,
                'desk': desk,
                'email': email,
                'section': current_section,
                'sheet': sheet_name,
                'section_head': None,  # Will be set when section ends
                'is_head': is_head
            }
            
            # Track section head
            if is_head and current_section_head is None:
                current_section_head = name
            
            section_staff.append(contact)
        
        # Finalize last section in this sheet
        for s in section_staff:
            s['section_head'] = current_section_head
        all_contacts.extend(section_staff)
    
    wb.close()
    _contacts_cache = all_contacts
    print(f"[Contacts] Loaded {len(all_contacts)} contacts from {len(wb.sheetnames)} sheets")
    return all_contacts

def get_contacts():
    global _contacts_cache
    if _contacts_cache is None:
        load_contacts_from_excel()
    return _contacts_cache

@app.route('/api/auth/login', methods=['POST'])
def auth_login():
    data = request.get_json() or {}
    email = data.get('email', '').strip().lower()
    
    if not email:
        return jsonify({"success": False, "message": "กรุณากรอกอีเมล"}), 400
    
    contacts = get_contacts()
    matched_contact = None
    
    # 1. Exact match on email field
    for c in contacts:
        if c.get('email') and str(c.get('email')).strip().lower() == email:
            matched_contact = c
            break
            
    # 2. Match username prefix (e.g., entering 'somchai' or 'somchai@gistda.or.th')
    if not matched_contact:
        username = email.split('@')[0]
        for c in contacts:
            c_email = str(c.get('email') or '').strip().lower()
            if c_email and (c_email == email or c_email.split('@')[0] == username):
                matched_contact = c
                break

    if matched_contact:
        return jsonify({
            "success": True,
            "user": {
                "name": matched_contact.get('name'),
                "nickname": matched_contact.get('nickname'),
                "position": matched_contact.get('position'),
                "email": matched_contact.get('email') or email,
                "section": matched_contact.get('section'),
                "sheet": matched_contact.get('sheet'),
                "is_head": matched_contact.get('is_head', False)
            }
        })
    
    # Fallback: If valid organizational domain, allow login as guest/staff
    if "@" in email:
        domain = email.split('@')[1]
        if domain == "gistda.or.th" or domain.endswith(".gistda.or.th"):
            fallback_user = {
                "name": email.split('@')[0],
                "nickname": None,
                "position": "บุคลากร สทอภ.",
                "email": email,
                "section": "สทอภ.",
                "sheet": "ทั่วไป",
                "is_head": False
            }
            return jsonify({
                "success": True,
                "user": fallback_user
            })

    return jsonify({
        "success": False,
        "message": "ไม่พบอีเมลนี้ในระบบบุคลากร กรุณาตรวจสอบอีเมลอีกครั้ง"
    }), 404


# Define Schema for Gemini Bill Extraction
class BillItem(BaseModel):
    item_code: Optional[str] = Field(None, description="The product code, SKU, or barcode (e.g. MS4-000939). If not available, leave null.")
    description: str = Field(description="Detailed item name or description in Thai or English as written on the bill. Keep serial numbers (S/N) if present.")
    quantity: int = Field(description="Quantity purchased")
    unit: str = Field("ชิ้น", description="The unit of measurement in Thai (e.g. 'ชิ้น', 'อัน', 'ลัง', 'เครื่อง', 'ถัง', 'กระป๋อง', 'ตัว') as appeared on the document. Default to 'ชิ้น' if not specified.")
    unit_price: float = Field(description="Unit price of the item")
    total_price: float = Field(description="Total price of the item (quantity * unit_price)")

class BillExtraction(BaseModel):
    vendor_name: str = Field(description="Name of the company or store that issued the bill/receipt (in Thai, e.g. 'บริษัท ไอที ซิตี้ (มหาชน)' or 'บริษัท บีทูเอส จำกัด')")
    invoice_number: str = Field(description="Invoice or Receipt number (เลขที่/เลขที่ใบเสร็จ/เล่มที่-เลขที่) as it appears on the document")
    invoice_date: str = Field(description="Date of the invoice/receipt in Thai format (e.g. '28 ตุลาคม 2567')")
    doc_type: str = Field("ใบเสร็จรับเงิน/ใบกำกับภาษี", description="Type of the document (e.g., 'ใบเสร็จรับเงิน', 'ใบกำกับภาษี', 'ใบเสร็จรับเงิน/ใบกำกับภาษี', 'บิลเงินสด', or other types as written on the header of the document. Default to 'ใบเสร็จรับเงิน/ใบกำกับภาษี' if not explicitly stated.)")
    items: List[BillItem] = Field(description="List of all purchased items")
    discount: float = Field(0.0, description="Total discount applied to the bill if any")
    grand_total: float = Field(description="Grand total / total amount paid in Baht")

class MemoAnalysis(BaseModel):
    intro_text: str = Field(description="A formal, polite Thai memo introduction paragraph starting with tab and '\\t\\tด้วย [หน่วยงาน] ได้ดำเนินการจัดซื้อวัสดุ...' explaining what was bought, summarizing key items naturally. Do not mention prices in this paragraph.")
    regulatory_text: str = Field(description="The most appropriate procurement regulation clause reference. Choose between: 1) หนังสือคณะกรรมการวินิจฉัยปัญหาการจัดซื้อจัดจ้างและการบริหารพัสดุภาครัฐ กรมบัญชีกลาง ด่วนที่สุด ที่ กค (กวจ) 0405.2/ว 119 ลงวันที่ 7 มีนาคม 2561 เรื่องแนวทางการปฎิบัติในการดำเนินการจัดหาพัสดุที่เกี่ยวกับค่าใช้จ่ายในการบริหารงาน ค่าใช้จ่ายในการฝึกอบรม การจัดงาน และการประชุมของหน่วยงานของรัฐ ตาราง 1 ลำดับที่ 3 if items are related to training/seminars/catering/activities. 2) ระเบียบกระทรวงการคลังว่าด้วยการจัดซื้อจัดจ้างและการบริหารพัสดุภาครัฐ พ.ศ. 2560 ตาราง 1 ลำดับที่ 1 if items are general office/computer/common supplies.")

# Font and formatting helpers for TH SarabunPSK
def set_run_font_enhanced(run, font_name='TH SarabunPSK', font_size_pt=16, bold=False, italic=False, underline=None, spacing=None):
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(font_size_pt)
    run.font.name = font_name
    
    if underline is not None:
        run.font.underline = underline
        
    rPr = run._r.get_or_add_rPr()
    
    # Set explicit font size in XML (sz and szCs in half-points)
    sz = rPr.find(qn('w:sz'))
    if sz is None:
        sz = OxmlElement('w:sz')
        rPr.append(sz)
    sz.set(qn('w:val'), str(int(font_size_pt * 2)))
    
    szCs = rPr.find(qn('w:szCs'))
    if szCs is None:
        szCs = OxmlElement('w:szCs')
        rPr.append(szCs)
    szCs.set(qn('w:val'), str(int(font_size_pt * 2)))
    
    if bold:
        # Set bold complex script
        bCs = rPr.find(qn('w:bCs'))
        if bCs is None:
            bCs = OxmlElement('w:bCs')
            rPr.append(bCs)
    if italic:
        # Set italic complex script
        iCs = rPr.find(qn('w:iCs'))
        if iCs is None:
            iCs = OxmlElement('w:iCs')
            rPr.append(iCs)
        
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rFonts.set(qn('w:cs'), font_name)
    rFonts.set(qn('w:eastAsia'), font_name)
    
    # Set character spacing (w:spacing) in dxa (dxa = 1/20 pt)
    if spacing is not None:
        spacing_elem = rPr.find(qn('w:spacing'))
        if spacing_elem is None:
            spacing_elem = OxmlElement('w:spacing')
            rPr.append(spacing_elem)
        spacing_elem.set(qn('w:val'), str(spacing))

def set_run_font(run, font_name='TH SarabunPSK', font_size_pt=16, bold=False, italic=False):
    set_run_font_enhanced(run, font_name, font_size_pt, bold, italic)

def copy_paragraph_format(src_p, dst_p):
    # Get the paragraph properties XML elements
    src_pPr = src_p._p.get_or_add_pPr()
    dst_pPr = dst_p._p.get_or_add_pPr()
    
    # Clear existing paragraph properties
    dst_pPr.clear()
    
    # Deep copy all child elements of paragraph properties (alignment, spacing, indents, styles, tab stops)
    for child in src_pPr:
        dst_pPr.append(copy.deepcopy(child))

def format_price(val):
    if val % 1 == 0:
        return f"{val:,.0f}"
    else:
        return f"{val:,.2f}"

def parse_thai_date(date_str):
    if not date_str:
        return (9999, 12, 31)
    
    date_str = str(date_str).strip()
    
    # Slash format: D/M/YYYY or D/M/B.E.
    slash_match = re.match(r'^(\d+)/(\d+)/(\d+)', date_str)
    if slash_match:
        d, m, y = map(int, slash_match.groups())
        if y > 2400:
            y -= 543
        return (y, m, d)
        
    # ISO Format: YYYY-MM-DD
    iso_match = re.match(r'^(\d{4})-(\d{2})-(\d{2})', date_str)
    if iso_match:
        y, m, d = map(int, iso_match.groups())
        if y > 2400:
            y -= 543
        return (y, m, d)
        
    THAI_MONTHS = {
        'มกราคม': 1, 'กุมภาพันธ์': 2, 'มีนาคม': 3, 'เมษายน': 4,
        'พฤษภาคม': 5, 'มิถุนายน': 6, 'กรกฎาคม': 7, 'สิงหาคม': 8,
        'กันยายน': 9, 'ตุลาคม': 10, 'พฤศจิกายน': 11, 'ธันวาคม': 12,
        'ม.ค.': 1, 'ก.พ.': 2, 'มี.ค.': 3, 'เม.ย.': 4,
        'พ.ค.': 5, 'มิ.ย.': 6, 'ก.ค.': 7, 'ส.ค.': 8,
        'ก.ย.': 9, 'ต.ค.': 10, 'พ.ย.': 11, 'ธ.ค.': 12
    }
    
    date_str = re.sub(r'\s+', ' ', date_str)
    parts = date_str.split(' ')
    
    day = 1
    month = 1
    year = 2026
    found_day = False
    
    for p in parts:
        p_clean = p.strip()
        if not p_clean:
            continue
            
        for m_name, m_val in THAI_MONTHS.items():
            if m_name in p_clean:
                month = m_val
                break
        else:
            if p_clean.isdigit():
                val = int(p_clean)
                if val > 2000:
                    year = val - 543 if val > 2400 else val
                else:
                    if not found_day:
                        day = val
                        found_day = True
                        
    return (year, month, day)


def update_paragraph_3(p, dept, phone):
    # Change department run
    for r in p.runs:
        if "สคร." in r.text:
            r.text = dept
            break
            
    # Change phone run
    phone_run_idx = None
    for idx, r in enumerate(p.runs):
        if "033 005 833" in r.text:
            phone_run_idx = idx
            break
            
    if phone_run_idx is not None:
        p.runs[phone_run_idx].text = phone
        # Clear the old "(ณัฐภัท)" runs instead of deleting them to preserve trailing tabs
        for i in range(phone_run_idx + 1, len(p.runs)):
            text_to_check = p.runs[i].text
            if any(char in text_to_check for char in ['(', 'ณัฐ', 'ภัท', ')']):
                p.runs[i].text = ""

def update_paragraph_4(p, memo_no, date_val):
    # Find the run containing "สคร."
    memo_start_idx = None
    for idx, r in enumerate(p.runs):
        if "สคร." in r.text:
            memo_start_idx = idx
            break
            
    # Find the run containing "วันที่"
    date_label_idx = None
    for idx, r in enumerate(p.runs):
        if "วันที่" in r.text:
            date_label_idx = idx
            break
            
    if memo_start_idx is not None and date_label_idx is not None:
        # Put the entire memo_no in the memo_start_idx run
        p.runs[memo_start_idx].text = memo_no
        # Clear the runs between memo_start_idx + 1 and date_label_idx - 1 (preserves spaces before "วันที่")
        for i in range(memo_start_idx + 1, date_label_idx - 1):
            p.runs[i].text = ""
            
    # Find the run containing "พฤศจิกายน"
    date_val_idx = None
    for idx, r in enumerate(p.runs):
        if "พฤศจิกายน" in r.text:
            date_val_idx = idx
            break
            
    if date_val_idx is not None:
        # Put the date value in the date_val_idx run
        p.runs[date_val_idx].text = f" {date_val} "
        # Clear the run containing "  2567" (which is typically date_val_idx + 1)
        if date_val_idx + 1 < len(p.runs) and "2567" in p.runs[date_val_idx + 1].text:
            p.runs[date_val_idx + 1].text = ""

def update_paragraph_5(p, subject):
    # Find "เรื่อง"
    subject_label_idx = None
    for idx, r in enumerate(p.runs):
        if "เรื่อง" in r.text:
            subject_label_idx = idx
            break
            
    if subject_label_idx is not None and subject_label_idx + 2 < len(p.runs):
        # Put the new subject text in Run 2 (subject_label_idx + 2)
        p.runs[subject_label_idx + 2].text = subject
        # Clear Runs 3 to 8 (which contained old subject portions)
        for i in range(subject_label_idx + 3, subject_label_idx + 9):
            if i < len(p.runs):
                p.runs[i].text = ""

def update_paragraph_6(p, to_text):
    # Find "เรียน"
    learn_idx = None
    for idx, r in enumerate(p.runs):
        if "เรียน" in r.text:
            learn_idx = idx
            break
            
    if learn_idx is not None and learn_idx + 2 < len(p.runs):
        # Put the new to_text in Run 2 (learn_idx + 2)
        p.runs[learn_idx + 2].text = to_text
        # Clear Runs 3 to 6
        for i in range(learn_idx + 3, len(p.runs)):
            p.runs[i].text = ""

def render_intro_paragraph(p, intro_text, dept, total_items):
    # Try parsing intro_text using regex to extract variables
    pattern = r"ด้วย\s+(.*?)\s+ได้ดำเนินการจัดซื้อวัสดุสำหรับการจัด\s+หลักสูตร\s+(.*?)\s+จำนวน\s+(.*?)\s+รายการ\s+โดยมีวัตถุประสงค์ตาม\s+(.*?)\s+โดยใช้งบประมาณ\s+\(รหัสงบประมาณ\s+(.*?)\s*ซึ่งมีรายละเอียดดังต่อไปนี้"
    match = re.search(pattern, intro_text, re.DOTALL)
    
    if match and len(p.runs) >= 41:
        val_dept = match.group(1).strip()
        val_course = match.group(2).strip()
        val_qty = match.group(3).strip()
        val_project = match.group(4).strip()
        val_budget = match.group(5).strip()
        
        # 1. Dept (Run 2)
        p.runs[2].text = val_dept
        
        # 2. Course (Run 4 and 5)
        p.runs[4].text = "ได้ดำเนินการจัดซื้อวัสดุสำหรับการจัด หลักสูตร " + val_course
        p.runs[5].text = ""
        
        # 3. Qty (Run 7)
        p.runs[7].text = str(total_items)
        
        # 4. Project (Run 12 to 20)
        p.runs[12].text = val_project
        for i in range(13, 21):
            p.runs[i].text = ""
            
        # 5. Budget (Run 24 to 38)
        p.runs[24].text = " " + val_budget
        for i in range(25, 39):
            p.runs[i].text = ""
        return
        
    # Fallback
    p.text = ""
    run = p.add_run(intro_text)
    set_run_font_enhanced(run, 'TH SarabunPSK', 16)

def update_regulatory_paragraph(p, reg_text):
    # Search for list index value at the end
    match = re.search(r"ตาราง\s+1\s+ลำดับที่\s+(.*)", reg_text)
    if match and len(p.runs) >= 13:
        idx_val = match.group(1).strip()
        # Set Run 12 to index value (which has underline)
        p.runs[12].text = f" {idx_val}"
        # Set Run 11 to empty space
        p.runs[11].text = ""
        return
        
    # Fallback: recreate paragraph runs manually
    p.text = ""
    run = p.add_run(reg_text)
    set_run_font_enhanced(run, 'TH SarabunPSK', 16)

def add_item_paragraph(doc, current_anchor, ref_item_p, idx, item):
    new_p = doc.add_paragraph()
    current_anchor._element.addnext(new_p._element)
    copy_paragraph_format(ref_item_p, new_p)
    
    # Override alignment to standard Justified to prevent weird stretching in thaiDistribute
    new_p.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY
    
    desc_val = item["desc"]
    if item["code"]:
        desc_val = f"{item['code']} {desc_val}"
        
    price_str = format_price(item["price"])
    
    # Precise run structures matching paragraph 8 but using regular spaces and no character spacing compression.
    runs_data = [
        ("\t", False, None),
        (f"{idx}. ค่า", False, None),
        ("\t", True, None),
        (desc_val, True, None),
        (" ", True, None), # Breaking space
        ("จำนวน", False, None),
        ("  ", True, None), # Regular spaces
        (str(item["qty"]), True, None),
        ("  ", True, None), # Regular spaces
        ("รายการ", False, None),
        (" ", False, None), # Breaking space
        ("เป็นเงิน", False, None),
        ("  ", True, None), # Breaking space
        (price_str, True, None),
        ("  ", True, None), # Regular spaces
        ("บาท", False, None),
        ("จากบริษัท", False, None),
        ("  ", True, None), # Breaking space
        (item["vendor"], True, None), # Normal spaces allowed to break naturally if long
        ("   ", True, None), # Breaking space
        ("ตามหลักฐานการจัดซื้อเป็น", False, None),
        (f"  {item['doc_type']}", True, None), # Breaking space before doc_type
        (" ", False, None),
        ("เล่มที่-เลขที่ ", False, None),
        (" ", True, None),
        (item["invoice_no"], True, None),
        ("  ", True, None),
        ("วันที่", False, None),
        (" ", True, None), # Regular space before date
        (item["invoice_date"] if item["invoice_date"] else "", True, None), # Date with regular spaces
        ("  ", True, None)
    ]
    
    for text_val, is_underlined, spacing_val in runs_data:
        run = new_p.add_run(text_val)
        underline_style = docx.enum.text.WD_UNDERLINE.DOTTED if is_underlined else False
        set_run_font_enhanced(run, 'TH SarabunPSK', 16, underline=underline_style, spacing=spacing_val)
        
    return new_p

def add_note_paragraph(doc, current_anchor, ref_note_p, start_num, end_num, subtotal, discount):
    new_p = doc.add_paragraph()
    current_anchor._element.addnext(new_p._element)
    copy_paragraph_format(ref_note_p, new_p)
    
    subtotal_str = format_price(subtotal)
    discount_str = format_price(discount)
    
    note_text = f"หมายเหตุ : รายการที่ {start_num}–{end_num} ยอดรวม {subtotal_str} บาท ได้รับส่วนลดทั้งหมด เป็นเงิน {discount_str} บาท    "
    
    run = new_p.add_run(note_text)
    set_run_font_enhanced(run, 'TH SarabunPSK', 16, underline=docx.enum.text.WD_UNDERLINE.DOTTED, spacing=None)
    
    return new_p

def update_total_paragraph(total_p, total_items, grand_total, thai_text):
    # Clear runs
    for r in list(total_p.runs):
        total_p._p.remove(r._r)
        
    grand_total_str = f"{grand_total:,.2f}"
    
    # New format: "                                         รวมเป็นจำนวนเงินทั้งสิ้น   7,618.00  บาท     "
    # All bold, no trailing Thai text
    runs_data = [
        ("                                         ", False, False, None),
        ("รวมเป็นจำนวนเงินทั้งสิ้น", True, False, None),
        ("   ", True, False, None),
        (grand_total_str, True, False, None),
        ("  บาท     ", True, False, None),
    ]
    
    for text_val, is_bold, is_underlined, spacing_val in runs_data:
        run = total_p.add_run(text_val)
        underline_style = docx.enum.text.WD_UNDERLINE.DOTTED if is_underlined else False
        set_run_font_enhanced(run, 'TH SarabunPSK', 16, bold=is_bold, underline=underline_style, spacing=spacing_val)

@app.route('/')
def index():
    return app.send_static_file('index.html')

@app.route('/api/contacts')
def api_contacts():
    """Return all contacts from the organization Excel file."""
    contacts = get_contacts()
    return jsonify(contacts)

@app.route('/api/contacts/search')
def api_contacts_search():
    """Search contacts by name substring (case-insensitive)."""
    q = request.args.get('q', '').strip()
    if not q or len(q) < 1:
        return jsonify([])
    
    contacts = get_contacts()
    results = [c for c in contacts if q.lower() in c['name'].lower()]
    return jsonify(results[:20])  # Limit to 20 results

@app.route('/api/extract-bill', methods=['POST'])
def api_extract_bill():
    """
    Endpoint for local backend or Hugging Face Spaces to process receipt images via PaddleOCR.
    Accepts multipart/form-data ('file') or application/json ({'image': base64_str}).
    Returns: { "words": [...], "rawText": "..." }
    """
    image_bytes = None

    if 'file' in request.files:
        file = request.files['file']
        if file.filename != '':
            image_bytes = file.read()
    elif request.is_json:
        data = request.get_json(silent=True) or {}
        img_data = data.get('image') or data.get('image_base64')
        if img_data:
            if ',' in img_data:
                img_data = img_data.split(',', 1)[1]
            try:
                image_bytes = base64.b64decode(img_data)
            except Exception as b64_err:
                return jsonify({"error": f"Invalid base64 image data: {str(b64_err)}"}), 400

    if not image_bytes:
        return jsonify({"error": "No image data provided. Provide a 'file' or JSON 'image' payload."}), 400

    try:
        from ocr_service import extract_bill_data
        result = extract_bill_data(image_bytes)
        return jsonify(result)
    except Exception as e:
        print(f"OCR Error: {str(e)}")
        return jsonify({"error": str(e)}), 500

@app.route('/api/validate_key', methods=['POST'])
def validate_key():
    data = request.json or {}
    provider = data.get('provider', 'gemini')
    api_key = data.get('api_key') or request.headers.get('Authorization')
    if api_key and api_key.startswith('Bearer '):
        api_key = api_key[len('Bearer '):]
    if not api_key:
        api_key = HARDCODED_API_KEY
        if api_key and api_key.startswith('gsk_'):
            provider = 'groq'
        elif api_key and api_key.startswith('sk-'):
            provider = 'openai'
    
    if not api_key:
        return jsonify({"valid": False, "error": "กรุณากรอก API Key"}), 400
        
    if provider == 'openai':
        try:
            # Validate OpenAI API key using the /v1/models endpoint
            headers = {"Authorization": f"Bearer {api_key}"}
            response = requests.get("https://api.openai.com/v1/models", headers=headers, timeout=10)
            if response.status_code == 200:
                return jsonify({
                    "valid": True, 
                    "message": "OpenAI API Key สามารถใช้งานได้ปกติ!"
                })
            else:
                try:
                    err_json = response.json()
                    err_msg = err_json.get("error", {}).get("message", "API Key ไม่ถูกต้อง")
                except:
                    err_msg = response.text or "API Key ไม่ถูกต้อง"
                return jsonify({
                    "valid": False,
                    "error": f"OpenAI API Key ไม่ถูกต้อง: {err_msg}",
                    "tip": "กรุณาตรวจสอบเครดิตและการสะกดคำของ API Key หรือสร้างคีย์ใหม่ที่ OpenAI Platform"
                })
        except Exception as e:
            return jsonify({
                "valid": False,
                "error": f"เชื่อมต่อไปยัง OpenAI ล้มเหลว: {str(e)}",
                "tip": "กรุณาตรวจสอบการเชื่อมต่ออินเทอร์เน็ตของเครื่องเซิร์ฟเวอร์"
            })
            
    if provider == 'groq':
        try:
            # Validate Groq API key using the /v1/models endpoint
            headers = {"Authorization": f"Bearer {api_key}"}
            response = requests.get("https://api.groq.com/openai/v1/models", headers=headers, timeout=10)
            if response.status_code == 200:
                return jsonify({
                    "valid": True, 
                    "message": "Groq API Key สามารถใช้งานได้ปกติ!"
                })
            else:
                try:
                    err_json = response.json()
                    err_msg = err_json.get("error", {}).get("message", "API Key ไม่ถูกต้อง")
                except:
                    err_msg = response.text or "API Key ไม่ถูกต้อง"
                return jsonify({
                    "valid": False,
                    "error": f"Groq API Key ไม่ถูกต้อง: {err_msg}",
                    "tip": "กรุณาตรวจสอบการสะกดคำของ API Key หรือสร้างคีย์ใหม่ที่ Groq Console (https://console.groq.com/)"
                })
        except Exception as e:
            return jsonify({
                "valid": False,
                "error": f"เชื่อมต่อไปยัง Groq ล้มเหลว: {str(e)}",
                "tip": "กรุณาตรวจสอบการเชื่อมต่ออินเทอร์เน็ตของเครื่องเซิร์ฟเวอร์"
            })
            
    # Default to Gemini
    try:
        client = genai.Client(api_key=api_key)
        
        # Test content generation with the primary models to ensure the key actually works
        try:
            client.models.generate_content(
                model='gemini-2.5-flash',
                contents='Hello'
            )
        except Exception as e25:
            # Fallback test with gemini-1.5-flash
            client.models.generate_content(
                model='gemini-1.5-flash',
                contents='Hello'
            )
            
        return jsonify({
            "valid": True, 
            "message": "API Key สามารถใช้งานได้ปกติ!"
        })
    except Exception as e:
        err_msg = str(e)
        import traceback
        traceback.print_exc()
        
        friendly_err = err_msg
        tip = "กรุณาสร้าง API Key ใหม่จาก Google AI Studio (https://aistudio.google.com/)"
        
        if "PERMISSION_DENIED" in err_msg or "denied" in err_msg.lower():
            friendly_err = "สิทธิ์การเข้าใช้งานถูกปฏิเสธ (403 Permission Denied) คีย์หรือโปรเจกต์ของคุณถูกระงับสิทธิ์เข้าใช้งานระบบจาก Google"
        elif "API_KEY_INVALID" in err_msg or "key is invalid" in err_msg.lower():
            friendly_err = "API Key ไม่ถูกต้อง (400 Invalid API Key) กรุณาตรวจสอบและคัดลอกคีย์ใหม่อีกครั้ง"
        elif "NOT_FOUND" in err_msg or "not found" in err_msg.lower():
            friendly_err = "ไม่พบบริการโมเดล (404 Not Found) บริการ Google Generative Language API ยังไม่เปิดใช้งานในโปรเจกต์ของคีย์นี้"
            tip = "หากใช้ GCP API Key กรุณาเข้าไปเปิดใช้งาน 'Generative Language API' ใน Google Cloud Console หรือแนะนำให้สร้างคีย์ฟรีจาก Google AI Studio แทน"
            
        return jsonify({
            "valid": False,
            "error": friendly_err,
            "tip": tip
        })


def call_gemini_with_retries(client, img, prompt, schema=None):
    models_to_try = ['gemini-2.5-flash', 'gemini-2.0-flash', 'gemini-1.5-flash']
    
    # Try with schema first if requested
    if schema:
        for model_name in models_to_try:
            for attempt in range(3):
                try:
                    print(f"Attempting extraction with {model_name} (with schema), attempt {attempt+1}...")
                    response = client.models.generate_content(
                        model=model_name,
                        contents=[img, prompt],
                        config=types.GenerateContentConfig(
                            response_mime_type="application/json",
                            response_schema=schema,
                        ),
                    )
                    return response.text
                except Exception as e:
                    err_msg = str(e)
                    print(f"{model_name} (with schema) attempt {attempt+1} failed: {err_msg}")
                    # If it's a 503 or 429, wait and retry
                    if any(code in err_msg for code in ["503", "429"]) or any(kw in err_msg.lower() for kw in ["temporarily", "limit", "demand", "busy", "unavailable"]):
                        time.sleep(2 ** attempt)  # 1s, 2s, 4s backoff
                    else:
                        break # Go to next model immediately for 404, 403, etc.
                        
    # Fallback to plain text JSON prompts
    plain_prompt = prompt + "\nRespond ONLY with a valid JSON matching this schema: {\"vendor_name\": \"...\", \"invoice_number\": \"...\", \"invoice_date\": \"...\", \"doc_type\": \"...\", \"items\": [{\"item_code\": \"...\", \"description\": \"...\", \"quantity\": 1, \"unit\": \"ชิ้น\", \"unit_price\": 1.0, \"total_price\": 1.0}], \"discount\": 0.0, \"grand_total\": 0.0}. Do not include markdown formatting or backticks."
    
    for model_name in models_to_try:
        for attempt in range(3):
            try:
                print(f"Attempting extraction with {model_name} (plain JSON), attempt {attempt+1}...")
                response = client.models.generate_content(
                    model=model_name,
                    contents=[img, plain_prompt]
                )
                text = response.text.strip()
                # Clean up markdown code block if present
                if text.startswith("```"):
                    lines = text.split("\n")
                    if lines[0].startswith("```"):
                        lines = lines[1:]
                    if lines[-1].startswith("```"):
                        lines = lines[:-1]
                    text = "\n".join(lines).strip()
                return text
            except Exception as e:
                err_msg = str(e)
                print(f"{model_name} (plain JSON) attempt {attempt+1} failed: {err_msg}")
                if any(code in err_msg for code in ["503", "429"]) or any(kw in err_msg.lower() for kw in ["temporarily", "limit", "demand", "busy", "unavailable"]):
                    time.sleep(2 ** attempt)
                else:
                    break
                    
def call_openai_with_retries(api_key, img_bytes, prompt, mime_type="image/jpeg"):
    # Base64 encode the image
    base64_image = base64.b64encode(img_bytes).decode('utf-8')
    
    headers = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type": "application/json"
    }
    
    # We ask for a JSON object matching the BillExtraction structure
    json_prompt = (
        prompt + "\n"
        "Return the output in JSON format matching this schema:\n"
        "{\n"
        "  \"vendor_name\": \"Name of the company/store (Thai/English)\",\n"
        "  \"invoice_number\": \"Invoice/receipt number\",\n"
        "  \"invoice_date\": \"Invoice date in Thai format (e.g. '28 ตุลาคม 2567')\",\n"
        "  \"doc_type\": \"Type of the document (e.g. 'ใบเสร็จรับเงิน', 'ใบกำกับภาษี', 'ใบเสร็จรับเงิน/ใบกำกับภาษี', 'บิลเงินสด', or other types as written on the header of the document. Default to 'ใบเสร็จรับเงิน/ใบกำกับภาษี' if not explicitly stated.)\",\n"
        "  \"items\": [\n"
        "    {\n"
        "      \"item_code\": \"Product code/SKU/barcode if available, else null\",\n"
        "      \"description\": \"Item name/details (keep serial numbers if present)\",\n"
        "      \"quantity\": 1,\n"
        "      \"unit\": \"The unit of measurement in Thai (e.g. 'ชิ้น', 'อัน', 'ลัง', 'เครื่อง', 'ถัง', 'กระป๋อง', 'ตัว') as appeared on the document. Default to 'ชิ้น' if not specified.\",\n"
        "      \"unit_price\": 100.0,\n"
        "      \"total_price\": 100.0\n"
        "    }\n"
        "  ],\n"
        "  \"discount\": 0.0,\n"
        "  \"grand_total\": 100.0\n"
        "}"
    )
    
    payload = {
        "model": "gpt-4o-mini",
        "response_format": { "type": "json_object" },
        "messages": [
            {
                "role": "user",
                "content": [
                    {
                        "type": "text",
                        "text": json_prompt
                    },
                    {
                        "type": "image_url",
                        "image_url": {
                            "url": f"data:{mime_type};base64,{base64_image}"
                        }
                    }
                ]
            }
        ],
        "max_tokens": 4096
    }
    
    for attempt in range(3):
        try:
            print(f"Attempting OpenAI extraction (gpt-4o-mini), attempt {attempt+1}...")
            response = requests.post(
                "https://api.openai.com/v1/chat/completions",
                headers=headers,
                json=payload,
                timeout=60
            )
            if response.status_code == 200:
                result_json = response.json()
                content = result_json["choices"][0]["message"]["content"]
                return content
            else:
                err_text = response.text
                print(f"OpenAI failed with status {response.status_code}: {err_text}")
                if response.status_code in [429, 500, 503]:
                    if attempt == 2:
                        raise Exception(f"OpenAI API Error: {err_text}")
                    time.sleep(2 ** attempt)
                else:
                    raise Exception(f"OpenAI API Error: {err_text}")
        except Exception as e:
            if attempt == 2:
                raise e
            time.sleep(2 ** attempt)
            
    raise Exception("OpenAI request timed out or failed.")

def call_groq_with_retries(api_key, img_bytes, prompt, mime_type="image/jpeg"):
    # Base64 encode the image
    base64_image = base64.b64encode(img_bytes).decode('utf-8')
    
    headers = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type": "application/json"
    }
    
    # We ask for a JSON object matching the BillExtraction structure
    json_prompt = (
        prompt + "\n"
        "Return the output in JSON format matching this schema:\n"
        "{\n"
        "  \"vendor_name\": \"Name of the company/store (Thai/English)\",\n"
        "  \"invoice_number\": \"Invoice/receipt number\",\n"
        "  \"invoice_date\": \"Invoice date in Thai format (e.g. '28 ตุลาคม 2567')\",\n"
        "  \"doc_type\": \"Type of the document (e.g. 'ใบเสร็จรับเงิน', 'ใบกำกับภาษี', 'ใบเสร็จรับเงิน/ใบกำกับภาษี', 'บิลเงินสด', or other types as written on the header of the document. Default to 'ใบเสร็จรับเงิน/ใบกำกับภาษี' if not explicitly stated.)\",\n"
        "  \"items\": [\n"
        "    {\n"
        "      \"item_code\": \"Product code/SKU/barcode if available, else null\",\n"
        "      \"description\": \"Item name/details (keep serial numbers if present)\",\n"
        "      \"quantity\": 1,\n"
        "      \"unit\": \"The unit of measurement in Thai (e.g. 'ชิ้น', 'อัน', 'ลัง', 'เครื่อง', 'ถัง', 'กระป๋อง', 'ตัว') as appeared on the document. Default to 'ชิ้น' if not specified.\",\n"
        "      \"unit_price\": 100.0,\n"
        "      \"total_price\": 100.0\n"
        "    }\n"
        "  ],\n"
        "  \"discount\": 0.0,\n"
        "  \"grand_total\": 100.0\n"
        "}"
    )
    
    payload = {
        "model": "meta-llama/llama-4-scout-17b-16e-instruct",
        "response_format": { "type": "json_object" },
        "messages": [
            {
                "role": "user",
                "content": [
                    {
                        "type": "text",
                        "text": json_prompt
                    },
                    {
                        "type": "image_url",
                        "image_url": {
                            "url": f"data:{mime_type};base64,{base64_image}"
                        }
                    }
                ]
            }
        ],
        "max_tokens": 4096
    }
    
    for attempt in range(3):
        try:
            print(f"Attempting Groq extraction (llama-3.2-11b-vision-preview), attempt {attempt+1}...")
            response = requests.post(
                "https://api.groq.com/openai/v1/chat/completions",
                headers=headers,
                json=payload,
                timeout=60
            )
            if response.status_code == 200:
                result_json = response.json()
                content = result_json["choices"][0]["message"]["content"]
                return content
            else:
                err_text = response.text
                print(f"Groq failed with status {response.status_code}: {err_text}")
                if response.status_code in [429, 500, 503]:
                    if attempt == 2:
                        raise Exception(f"Groq API Error: {err_text}")
                    time.sleep(2 ** attempt)
                else:
                    raise Exception(f"Groq API Error: {err_text}")
        except Exception as e:
            if attempt == 2:
                raise e
            time.sleep(2 ** attempt)
            
    raise Exception("Groq request timed out or failed.")

@app.route('/api/extract', methods=['POST'])
def extract_bill():
    provider = request.form.get('provider', 'gemini')
    api_key = request.headers.get('Authorization')
    if api_key and api_key.startswith('Bearer '):
        api_key = api_key[len('Bearer '):]
    
    if not api_key:
        api_key = request.form.get('api_key')
        
    if not api_key:
        api_key = HARDCODED_API_KEY
        if api_key and api_key.startswith('gsk_'):
            provider = 'groq'
        elif api_key and api_key.startswith('sk-'):
            provider = 'openai'
            
    if not api_key:
        return jsonify({"error": "API Key is required"}), 400
        
    if 'file' not in request.files:
        return jsonify({"error": "No file uploaded"}), 400
        
    file = request.files['file']
    if file.filename == '':
        return jsonify({"error": "No file selected"}), 400
        
    mime_type = file.content_type or "image/jpeg"
    
    prompt = """
    คุณคือผู้เชี่ยวชาญด้านการถอดข้อมูลบิลและใบเสร็จ (OCR & Receipt Parsing Specialist)
    หน้าที่ของคุณคือการวิเคราะห์ภาพถ่ายบิล/ใบเสร็จนี้อย่างแม่นยำ และแปลงข้อมูลเป็น JSON ตามโครงสร้างที่กำหนด
    
    คำแนะนำที่ต้องปฏิบัติตามอย่างเคร่งครัด:
    1. **ชื่อร้านค้า/บริษัทผู้ขาย (vendor_name)**: 
       - ต้องเป็นชื่อของห้างร้านหรือบริษัทที่ออกเอกสารนี้ (เช่น "บริษัท ไอที ซิตี้ (มหาชน)", "บริษัท ซีโอแอล จำกัด (มหาชน)") ซึ่งมักจะอยู่ด้านบนสุดของบิล
       - ห้ามสับสนกับชื่อผู้ซื้อ (เช่น "สำนักงาน...", "สคร.") หรือชื่อสินค้าเด็ดขาด!
       - หากชื่อร้านเป็นภาษาไทย ให้ใช้ภาษาไทย
    2. **เลขที่เอกสาร/ใบเสร็จ (invoice_number)**:
       - มองหาคำว่า "เลขที่", "เลขที่ใบเสร็จ", "เลขที่ใบกำกับภาษี", "Invoice No.", "Receipt No.", "Tax Invoice No.", "Doc No."
       - ห้ามใช้เลขประจำตัวผู้เสียภาษี (Tax ID ซึ่งเป็นเลข 13 หลัก) หรือเลขที่สาขา มาเป็นเลขที่เอกสารเด็ดขาด!
    3. **วันที่เอกสาร (invoice_date)**:
       - ต้องแปลงเป็นฟอร์แมตภาษาไทย เช่น "28 ตุลาคม 2567" หรือ "9 มิถุนายน 2569" 
       - หากบิลระบุปีคริสตศักราช (ค.ศ. เช่น 2024, 2025, 2026) ต้องแปลงเป็นปีพุทธศักราช (พ.ศ. โดยบวก 543 เช่น 2024 -> 2567, 2026 -> 2569)
    4. **ประเภทเอกสาร (doc_type)**:
       - ตรวจสอบและระบุประเภทเอกสารตามที่ปรากฏบนหัวเอกสาร เช่น "ใบเสร็จรับเงิน", "ใบกำกับภาษี", "ใบเสร็จรับเงิน/ใบกำกับภาษี", "บิลเงินสด" หรือคำระบุประเภทเอกสารอื่นๆ
       - หากไม่แน่ใจหรือไม่มีระบุไว้ให้ใช้ "ใบเสร็จรับเงิน/ใบกำกับภาษี" เป็นค่าเริ่มต้น
    5. **รายการสินค้า (items)**:
       - ดึงรายการสินค้าทั้งหมดออกมาในรูปแบบอาร์เรย์
       - **รหัสสินค้า (item_code)**: หากมีรหัสสินค้า (SKU, Barcode, Product Code เช่น MS4-000939 หรือเลขบาร์โค้ด) ให้ดึงมาใส่ในฟิลด์นี้ หากไม่มีให้ใส่ null
       - **รายละเอียดสินค้า (description)**: ต้องดึงชื่อสินค้าและรายละเอียดพัสดุมาให้ครบถ้วนทุกตัวอักษร **ห้ามย่อ ตัดทอน หรือสรุปชื่อสินค้าเด็ดขาด (ห้ามตัดคำ)** ให้คัดลอกข้อความยาวๆ ทั้งหมดแบบเป๊ะๆ 100% ตามที่ปรากฏในภาพ
       - **จำนวน (quantity)**: ต้องเป็นจำนวนชิ้น (จำนวนเต็ม)
       - **ราคาต่อหน่วย (unit_price)**: ราคาต่อหน่วยก่อนหักส่วนลด (ทศนิยม)
       - **ราคารวม (total_price)**: ราคารวมของรายการนั้น (quantity * unit_price)
       - **ความถูกต้องของตัวเลข**: ตรวจสอบคำนวณทางคณิตศาสตร์ให้ถูกต้องเสมอ!
    6. **ส่วนลด (discount)**:
       - ยอดส่วนลดรวมท้ายบิล (ถ้ามี) หากไม่มีให้ใส่ 0.0
    7. **ยอดเงินรวมทั้งสิ้น (grand_total)**:
       - ยอดเงินสุทธิรวมที่ต้องจ่ายจริงท้ายบิล ตรวจสอบให้มั่นใจว่าตัวเลขราคาและยอดรวมถูกต้องตรงตามใบเสร็จ
    """
    
    try:
        if provider == 'openai':
            # Read file bytes for base64 encoding
            file.stream.seek(0)
            img_bytes = file.stream.read()
            response_text = call_openai_with_retries(api_key, img_bytes, prompt, mime_type=mime_type)
            return response_text, 200, {'Content-Type': 'application/json'}
        elif provider == 'groq':
            # Read file bytes for base64 encoding
            file.stream.seek(0)
            img_bytes = file.stream.read()
            response_text = call_groq_with_retries(api_key, img_bytes, prompt, mime_type=mime_type)
            return response_text, 200, {'Content-Type': 'application/json'}
        else:
            # Default to Gemini
            img = Image.open(file.stream)
            client = genai.Client(api_key=api_key)
            response_text = call_gemini_with_retries(client, img, prompt, schema=BillExtraction)
            return response_text, 200, {'Content-Type': 'application/json'}
        
    except Exception as e:
        import traceback
        traceback.print_exc()
        
        err_msg = str(e)
        tip_msg = ""
        if provider == 'gemini':
            if "NOT_FOUND" in err_msg or "not found" in err_msg.lower() or "404" in err_msg:
                tip_msg = "\n\n💡 Tip: This usually means the Generative Language API is not enabled for your project, or your API key is invalid. Please create a key from Google AI Studio (https://aistudio.google.com/) and try again."
        elif provider == 'openai':
            tip_msg = "\n\n💡 Tip: Please verify that your OpenAI API key is correct, has credits available, and the OpenAI service is online."
        elif provider == 'groq':
            tip_msg = "\n\n💡 Tip: Please verify that your Groq API key is correct and you have not exceeded your Groq rate limits. Check console.groq.com."
            
        return jsonify({"error": f"Failed during bill extraction: {err_msg}{tip_msg}"}), 500

@app.route('/api/analyze_purchase', methods=['POST'])
def analyze_purchase():
    data = request.json or {}
    provider = data.get('provider', 'gemini')
    api_key = data.get('api_key') or request.headers.get('Authorization')
    if api_key and api_key.startswith('Bearer '):
        api_key = api_key[len('Bearer '):]
    if not api_key:
        api_key = HARDCODED_API_KEY
        if api_key and api_key.startswith('gsk_'):
            provider = 'groq'
        elif api_key and api_key.startswith('sk-'):
            provider = 'openai'
    
    if not api_key:
        return jsonify({"error": "API Key is required"}), 400
        
    items = data.get('items', [])
    department = data.get('department', 'สคร.')
    
    if not items:
        return jsonify({"error": "No items provided for analysis"}), 400
        
    items_list_str = "\n".join([f"- {item.get('description', '')} (จำนวน {item.get('quantity', 1)})" for item in items])
    
    prompt = f"""
    Analyze the following list of items purchased by the department '{department}':
    {items_list_str}
    
    Task:
    1. Write a professional, formal Thai memo introduction paragraph starting with two tabs: '\\t\\tด้วย {department} ได้ดำเนินการจัดซื้อวัสดุสำหรับการจัด ...' or 'ด้วย {department} ได้ดำเนินการจัดซื้อวัสดุ ...'. Describe what was bought by summarizing the key items naturally and politely. Do not mention total prices in this paragraph.
    2. Suggest the most appropriate Thai procurement regulation clause from these choices:
       - Choose "หนังสือคณะกรรมการวินิจฉัยปัญหาการจัดซื้อจัดจ้างและการบริหารพัสดุภาครัฐ กรมบัญชีกลาง ด่วนที่สุด ที่ กค (กวจ) 0405.2/ว 119 ลงวันที่ 7 มีนาคม 2561 เรื่องแนวทางการปฎิบัติในการดำเนินการจัดหาพัสดุที่เกี่ยวกับค่าใช้จ่ายในการบริหารงาน ค่าใช้จ่ายในการฝึกอบรม การจัดงาน และการประชุมของหน่วยงานของรัฐ ตาราง 1 ลำดับที่ 3" if the items are for training, seminars, workshops, meetings, catering, or similar educational activities.
       - Choose "ระเบียบกระทรวงการคลังว่าด้วยการจัดซื้อจัดจ้างและการบริหารพัสดุภาครัฐ พ.ศ. 2560 ตาราง 1 ลำดับที่ 1" if the items are general office supplies, computer equipment, network materials, hardware, or standard utility items.
    """
    
    try:
        if provider == 'openai' or provider == 'groq':
            # Call OpenAI or Groq (since their completions format is identical)
            headers = {
                "Authorization": f"Bearer {api_key}",
                "Content-Type": "application/json"
            }
            json_prompt = (
                prompt + "\n"
                "Return the output in JSON format matching this schema:\n"
                "{\n"
                "  \"intro_text\": \"Generated intro paragraph...\",\n"
                "  \"regulatory_text\": \"Selected regulation text...\"\n"
                "}"
            )
            payload = {
                "model": "gpt-4o-mini" if provider == 'openai' else "llama-3.3-70b-versatile",
                "response_format": { "type": "json_object" },
                "messages": [
                    {
                        "role": "user",
                        "content": json_prompt
                    }
                ],
                "max_tokens": 1024
            }
            api_url = "https://api.openai.com/v1/chat/completions" if provider == 'openai' else "https://api.groq.com/openai/v1/chat/completions"
            response = requests.post(
                api_url,
                headers=headers,
                json=payload,
                timeout=30
            )
            if response.status_code == 200:
                result_json = response.json()
                content = result_json["choices"][0]["message"]["content"]
                return content, 200, {'Content-Type': 'application/json'}
            else:
                raise Exception(f"{provider.capitalize()} error: {response.text}")
                
        else:
            # Default to Gemini
            client = genai.Client(api_key=api_key)
            models_to_try = ['gemini-2.5-flash', 'gemini-2.0-flash', 'gemini-1.5-flash']
            
            for model_name in models_to_try:
                try:
                    response = client.models.generate_content(
                        model=model_name,
                        contents=prompt,
                        config=types.GenerateContentConfig(
                            response_mime_type="application/json",
                            response_schema=MemoAnalysis,
                        ),
                    )
                    return response.text, 200, {'Content-Type': 'application/json'}
                except Exception as e:
                    print(f"Gemini {model_name} analysis failed: {str(e)}")
                    continue
            raise Exception("All Gemini models failed to analyze the purchase.")
            
    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({"error": f"Failed to analyze purchase: {str(e)}"}), 500

@app.route('/api/generate', methods=['POST'])
def generate_docx():
    data = request.json
    if not data:
        return jsonify({"error": "Missing payload"}), 400
        
    template_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'data', 'แบบฟอร์ม รายงานขอความเห็นชอบซื้อจ้าง.docx')
    if not os.path.exists(template_path):
        return jsonify({"error": "Template document file not found on server"}), 500
        
    try:
        doc = docx.Document(template_path)
        
        # Process items to list first to get total_items_count
        invoices = data.get('invoices', [])
        # Sort invoices chronologically by date
        invoices = sorted(invoices, key=lambda x: parse_thai_date(x.get('invoice_date')))
        flat_items = []
        item_counter = 1
        invoice_ranges = []
        
        for inv in invoices:
            items_list = inv.get('items', [])
            if not items_list:
                continue
            
            # Merge duplicate items within this invoice (same code and description)
            merged_items = []
            for item in items_list:
                code = (item.get('item_code') or '').strip()
                desc = (item.get('description') or '').strip()
                qty = int(item.get('quantity', 1))
                unit_p = float(item.get('unit_price', 0.0))
                tot_p = float(item.get('total_price', 0.0))
                
                found = False
                for m in merged_items:
                    if m['code_key'].lower() == code.lower() and m['desc'].lower() == desc.lower():
                        m['qty'] += qty
                        m['price'] += tot_p
                        found = True
                        break
                if not found:
                    merged_items.append({
                        "code_key": code,
                        "code": code if code else None,
                        "desc": desc,
                        "qty": qty,
                        "price": tot_p
                    })
            
            start_num = item_counter
            for item in merged_items:
                flat_items.append({
                    "code": item["code"],
                    "desc": item["desc"],
                    "qty": item["qty"],
                    "price": item["price"],
                    "vendor": inv.get('vendor_name'),
                    "invoice_no": inv.get('invoice_number'),
                    "invoice_date": inv.get('invoice_date'),
                    "doc_type": inv.get('doc_type', 'ใบเสร็จรับเงิน/ใบกำกับภาษี')
                })
                item_counter += 1
            end_num = item_counter - 1
            
            invoice_ranges.append({
                "start": start_num,
                "end": end_num,
                "discount": inv.get('discount', 0.0),
                "subtotal": sum(item["price"] for item in merged_items),
                "vendor": inv.get('vendor_name')
            })
            
        total_items_count = len(flat_items)
        
        # 1. Update Header Paragraphs
        # Paragraph 3: Department & Phone
        update_paragraph_3(doc.paragraphs[3], data.get('department', 'สคร.'), data.get('phone', ''))
        
        # Paragraph 4: Memo No & Date
        # Clean date values to prevent wrapping
        doc_date = data.get('date', 'พฤศจิกายน  2567')
        update_paragraph_4(doc.paragraphs[4], data.get('memo_no', 'สคร.             /2567'), doc_date)
        
        # Paragraph 5: Subject
        subject_text = f"รายงานขอความเห็นชอบการจัดซื้อจัดจ้าง  จำนวน {total_items_count}  รายการ"
        update_paragraph_5(doc.paragraphs[5], data.get('subject', subject_text))
        
        # Paragraph 6: Learn (เรียน)
        update_paragraph_6(doc.paragraphs[6], data.get('to_text', 'ผอ.สคร.  ผ่าน รก.หน.ฝถท.'))
        
        # Paragraph 7: Introduction Body
        render_intro_paragraph(doc.paragraphs[7], data.get('intro_text', ''), data.get('department', 'สคร.'), total_items_count)
        
        # 2. Find and Replace the Items List
        start_idx = None
        end_idx = None
        
        for i, p in enumerate(doc.paragraphs):
            text = p.text.strip()
            if text.startswith("1. ค่า"):
                start_idx = i
                break
                
        for i, p in enumerate(doc.paragraphs):
            text = p.text.strip()
            if text.startswith("รวม") and "รายการ เป็นเงินทั้งสิ้น" in text:
                end_idx = i
                break
                
        if start_idx is None or end_idx is None:
            return jsonify({"error": "Invalid template format: Could not find items section"}), 500
            
        ref_item_p = doc.paragraphs[start_idx]
        ref_note_p = doc.paragraphs[start_idx + 1] if "หมายเหตุ" in doc.paragraphs[start_idx + 1].text else ref_item_p
        
        # Delete old items list
        paragraphs_to_delete = doc.paragraphs[start_idx:end_idx]
        for p in paragraphs_to_delete:
            p_elem = p._element
            p_elem.getparent().remove(p_elem)
            p._element = None
            p._p = None
            
        # Re-fetch anchor paragraph
        anchor_p = doc.paragraphs[start_idx - 1]
        
        # Generate new paragraphs
        current_anchor = anchor_p
        
        # Track which invoice range a flat item belongs to so we can insert the correct note
        for idx, item in enumerate(flat_items, 1):
            # 1. Add Item Paragraph
            new_item_p = add_item_paragraph(doc, current_anchor, ref_item_p, idx, item)
            current_anchor = new_item_p
            
            # 2. Find if this item belongs to an invoice that has a discount
            matching_range = None
            for r in invoice_ranges:
                if r["start"] <= idx <= r["end"]:
                    matching_range = r
                    break
                    
            if matching_range and matching_range["discount"] > 0 and idx == matching_range["end"]:
                # Add note paragraph
                new_note_p = add_note_paragraph(doc, current_anchor, ref_note_p, matching_range['start'], matching_range['end'], matching_range['subtotal'], matching_range['discount'])
                current_anchor = new_note_p
                
        # 3. Update Total Paragraph (It moved down, let's find the new index)
        total_p_idx = None
        for i, p in enumerate(doc.paragraphs):
            if p.text.strip().startswith("รวม") and "รายการ เป็นเงินทั้งสิ้น" in p.text:
                total_p_idx = i
                break
                
        if total_p_idx is not None:
            total_p = doc.paragraphs[total_p_idx]
            
            # Calculate grand total sums
            grand_subtotal = sum(item["price"] for item in flat_items)
            total_discount = sum(r["discount"] for r in invoice_ranges)
            grand_total_paid = grand_subtotal - total_discount
            
            thai_text_amount = bahttext(grand_total_paid)
            
            # Update total paragraph
            update_total_paragraph(total_p, total_items_count, grand_total_paid, thai_text_amount)
            
        # 4. Update regulatory note
        reg_p_idx = total_p_idx + 1 if total_p_idx is not None else None
        if reg_p_idx is not None and reg_p_idx < len(doc.paragraphs):
            reg_p = doc.paragraphs[reg_p_idx]
            update_regulatory_paragraph(reg_p, data.get('regulatory_text', ''))
            
        # 5. Update signature blocks
        # Replace names and positions in the remaining paragraphs
        # Clean date values to prevent wrapping
        req_date = data.get('requester_date', '   / พฤศจิกายน / 2567')
        app_date = data.get('approver_date', '   / พฤศจิกายน / 2567')
        
        for p in doc.paragraphs[reg_p_idx + 1:]:
            p_text = p.text
            
            # 5.1 Requester Position (Paragraph 61)
            if "เจ้าหน้าที่/ผู้รับผิดชอบ" in p_text:
                for r in p.runs:
                    if "เจ้าหน้าที่/ผู้รับผิดชอบ" in r.text:
                        r.text = data.get('requester_position', 'เจ้าหน้าที่/ผู้รับผิดชอบ')
                        break
                        
            # 5.2 Requester Name and Date (Paragraph 62)
            elif "ศิริพักตร์" in p_text or "เสมียนคิด" in p_text:
                # Replace name
                for r in p.runs:
                    if "ศิริพักตร์" in r.text or "เสมียนคิด" in r.text:
                        r.text = data.get('requester_name', 'นางสาวศิริพักตร์  เสมียนคิด')
                        break
                # Replace date
                date_start_idx = None
                for idx, r in enumerate(p.runs):
                    if "/" in r.text or "พฤศจิกายน" in r.text:
                        date_start_idx = idx
                        break
                if date_start_idx is not None:
                    # Put new date in the first date run
                    p.runs[date_start_idx].text = req_date
                    # Clear subsequent runs in-place to preserve trailing tabs
                    for i in range(date_start_idx + 1, date_start_idx + 6):
                        if i < len(p.runs):
                            p.runs[i].text = ""
                    
            # 5.3 Approver Position (Paragraph 66)
            elif "ผอ.สคร." in p_text:
                for r in p.runs:
                    if "ผอ.สคร." in r.text:
                        r.text = data.get('approver_position', 'ผอ.สคร.')
                        break
                        
            # 5.4 Approver Name (Paragraph 67)
            elif "ปราณปริยา" in p_text or "วงค์ษา" in p_text:
                for r in p.runs:
                    if "ปราณปริยา" in r.text or "วงค์ษา" in r.text:
                        r.text = data.get('approver_name', 'นางสาวปราณปริยา   วงค์ษา')
                        break
                        
            # 5.5 Approver Date (Paragraph 68)
            elif ("/ พฤศจิกายน" in p_text or "/" in p_text) and not ("ศิริพักตร์" in p_text or "เสมียนคิด" in p_text):
                date_start_idx = None
                for idx, r in enumerate(p.runs):
                    if "/" in r.text or "พฤศจิกายน" in r.text:
                        date_start_idx = idx
                        break
                if date_start_idx is not None:
                    # Put new date in the first date run
                    p.runs[date_start_idx].text = app_date
                    # Clear subsequent runs in-place to preserve trailing tabs
                    for i in range(date_start_idx + 1, date_start_idx + 3):
                        if i < len(p.runs):
                            p.runs[i].text = ""
                    
        # Save to temporary file and return
        temp_dir = tempfile.gettempdir()
        out_filename = "รายงานขอความเห็นชอบการจัดซื้อจัดจ้าง.docx"
        out_path = os.path.join(temp_dir, out_filename)
        
        doc.save(out_path)
        
        return send_file(out_path, as_attachment=True, download_name=out_filename)
        
    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({"error": f"Failed during document generation: {str(e)}"}), 500


@app.route('/api/generate_excel', methods=['POST'])
def generate_excel():
    data = request.json or {}
    is_illus = data.get('is_illustration', False)
    try:
        # Load the workbook
        src_file = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'data', 'สรุปค่าใช้จ่าย_เบิกเงินค่าพัสดุ.xlsx')
        if not os.path.exists(src_file):
            return jsonify({"error": "Template file 'สรุปค่าใช้จ่าย_เบิกเงินค่าพัสดุ.xlsx' not found."}), 404
            
        wb = openpyxl.load_workbook(src_file, data_only=False)
        sheet2 = wb.worksheets[1]
        
        # ---------------------------------------------
        # Generate Illustration Excel (is_illustration = True)
        # ---------------------------------------------
        if is_illus:
            # 1. Cache styles from template before clearing
            title_font = copy.copy(sheet2.cell(row=1, column=1).font)
            title_align = copy.copy(sheet2.cell(row=1, column=1).alignment)
            
            # Row 4 is header: A4 is 'ลำดับ' style, B4 is 'รายละเอียด' style, U4 is 'หมายเหตุ' style
            header_col1_font = copy.copy(sheet2.cell(row=4, column=1).font)
            header_col1_fill = copy.copy(sheet2.cell(row=4, column=1).fill)
            header_col1_align = copy.copy(sheet2.cell(row=4, column=1).alignment)
            
            header_col2_font = copy.copy(sheet2.cell(row=4, column=2).font)
            header_col2_fill = copy.copy(sheet2.cell(row=4, column=2).fill)
            header_col2_align = copy.copy(sheet2.cell(row=4, column=2).alignment)
            
            header_colU_font = copy.copy(sheet2.cell(row=4, column=21).font)
            header_colU_fill = copy.copy(sheet2.cell(row=4, column=21).fill)
            header_colU_align = copy.copy(sheet2.cell(row=4, column=21).alignment)
            
            # Row 5 and 6 of original template are vendor header and item style cells
            vendor_style_cells = [sheet2.cell(row=5, column=c) for c in range(1, 22)]
            item_style_cells = [sheet2.cell(row=6, column=c) for c in range(1, 22)]
            
            vendor_row_height = sheet2.row_dimensions[5].height or 52.2
            item_row_height = sheet2.row_dimensions[6].height or 36.0
            
            # Clear all merged cell ranges
            for merged_range in list(sheet2.merged_cells.ranges):
                sheet2.merged_cells.remove(merged_range)
                
            # Clear all rows from row 1 to max_row
            sheet2.delete_rows(1, sheet2.max_row)
            
            # Set target column width for Column P
            sheet2.column_dimensions['P'].width = 56.25
            sheet2.column_dimensions['N'].width = 15.0
            
            # Set page setup and print options
            sheet2.page_setup.orientation = sheet2.ORIENTATION_PORTRAIT
            sheet2.page_setup.paperSize = sheet2.PAPERSIZE_A4
            sheet2.page_setup.fitToWidth = 1
            sheet2.page_setup.fitToHeight = 0
            sheet2.sheet_properties.pageSetUpPr.fitToPage = True
            sheet2.print_options.gridLines = True
            if sheet2.views.sheetView:
                for view in sheet2.views.sheetView:
                    view.showGridLines = True
            else:
                sheet2.views.sheetView.append(openpyxl.worksheet.views.SheetView(showGridLines=True))
            
            # 2. Write Row 1: 'ภาพประกอบ' title
            c_title = sheet2.cell(row=1, column=1, value='ภาพประกอบ')
            c_title.font = title_font
            c_title.alignment = title_align
            sheet2.row_dimensions[1].height = 35.0
            sheet2.merge_cells(start_row=1, start_column=1, end_row=1, end_column=16)
            
            # 3. Write Row 2: Headers ('ลำดับ', 'รายละเอียด', 'ภาพ')
            c_h1 = sheet2.cell(row=2, column=1, value='ลำดับ')
            c_h1.font = header_col1_font
            c_h1.fill = header_col1_fill
            c_h1.alignment = header_col1_align
            
            c_h2 = sheet2.cell(row=2, column=2, value='รายละเอียด')
            c_h2.font = header_col2_font
            c_h2.fill = header_col2_fill
            c_h2.alignment = header_col2_align
            
            c_h3 = sheet2.cell(row=2, column=16, value='ภาพ')
            c_h3.font = header_colU_font
            c_h3.fill = header_colU_fill
            c_h3.alignment = header_colU_align
            
            # Apply header font and styling across the merged range
            for col in range(2, 16):
                cell = sheet2.cell(row=2, column=col)
                cell.font = header_col2_font
                cell.fill = header_col2_fill
                cell.alignment = header_col2_align
                
            sheet2.merge_cells(start_row=2, start_column=2, end_row=2, end_column=15)
            sheet2.row_dimensions[2].height = 74.25
            
            # Draw header borders
            from openpyxl.styles import Border, Side
            thin_side = Side(border_style="thin", color="000000")
            medium_side = Side(border_style="medium", color="000000")
            
            sheet2.cell(row=2, column=1).border = Border(left=thin_side, right=thin_side, bottom=medium_side)
            for col in range(2, 15):
                sheet2.cell(row=2, column=col).border = Border(bottom=medium_side)
            sheet2.cell(row=2, column=15).border = Border(right=thin_side, bottom=medium_side)
            sheet2.cell(row=2, column=16).border = Border(left=thin_side, right=thin_side, bottom=medium_side)
            
            # 4. Write data rows
            current_row = 3
            invoices = data.get('invoices', [])
            invoices = sorted(invoices, key=lambda x: parse_thai_date(x.get('invoice_date')))
            
            temp_files = []
            
            def apply_style_illus(sheet, row_idx, style_cells):
                for col_idx, src_cell in enumerate(style_cells, 1):
                    dst_cell = sheet.cell(row=row_idx, column=col_idx)
                    if src_cell.font: dst_cell.font = copy.copy(src_cell.font)
                    if src_cell.fill: dst_cell.fill = copy.copy(src_cell.fill)
                    if src_cell.alignment: dst_cell.alignment = copy.copy(src_cell.alignment)
                    if src_cell.number_format: dst_cell.number_format = src_cell.number_format
            
            for inv_idx, inv in enumerate(invoices, 1):
                items = inv.get('items', [])
                
                # Vendor Header Row
                sheet2.cell(row=current_row, column=1, value=inv_idx)
                sheet2.cell(row=current_row, column=2, value=inv.get('vendor_name', ''))
                
                apply_style_illus(sheet2, current_row, vendor_style_cells)
                for c in [15, 16]:
                    src_c = vendor_style_cells[c-1] if c-1 < len(vendor_style_cells) else vendor_style_cells[-1]
                    dst_c = sheet2.cell(row=current_row, column=c)
                    if src_c.font: dst_c.font = copy.copy(src_c.font)
                    if src_c.fill: dst_c.fill = copy.copy(src_c.fill)
                    if src_c.alignment: dst_c.alignment = copy.copy(src_c.alignment)
                
                sheet2.row_dimensions[current_row].height = vendor_row_height
                sheet2.merge_cells(start_row=current_row, start_column=2, end_row=current_row, end_column=14)
                
                sheet2.cell(row=current_row, column=1).border = Border(left=thin_side, right=thin_side)
                sheet2.cell(row=current_row, column=2).border = Border(left=thin_side)
                sheet2.cell(row=current_row, column=15).border = Border(right=thin_side)
                sheet2.cell(row=current_row, column=16).border = Border(left=thin_side, right=thin_side)
                
                current_row += 1
                
                # Item Rows
                for item_idx, item in enumerate(items, 1):
                    item_row = current_row
                    
                    code = item.get('item_code', '').strip()
                    desc = item.get('description', '').strip()
                    desc_val = f"{code} {desc}" if code else desc
                    
                    sheet2.cell(row=item_row, column=1, value=f"{inv_idx}.{item_idx}")
                    sheet2.cell(row=item_row, column=2, value=desc_val)
                    sheet2.cell(row=item_row, column=4, value="(")
                    sheet2.cell(row=item_row, column=5, value=float(item.get('unit_price', 0.0)))
                    sheet2.cell(row=item_row, column=6, value="บาท * ")
                    sheet2.cell(row=item_row, column=7, value=int(item.get('quantity', 1)))
                    sheet2.cell(row=item_row, column=8, value=item.get('unit', 'ชิ้น'))
                    sheet2.cell(row=item_row, column=9, value=")")
                    sheet2.cell(row=item_row, column=13, value="=")
                    sheet2.cell(row=item_row, column=14, value=f"=E{item_row}*G{item_row}")
                    sheet2.cell(row=item_row, column=15, value="บาท")
                    
                    apply_style_illus(sheet2, item_row, item_style_cells)
                    from openpyxl.styles import Alignment
                    desc_c = sheet2.cell(row=item_row, column=2)
                    desc_c.alignment = Alignment(wrapText=True, vertical='top', horizontal='left')
                    sheet2.merge_cells(start_row=item_row, start_column=2, end_row=item_row, end_column=3)
                    # Apply styles to column P
                    src_c = item_style_cells[15] if 15 < len(item_style_cells) else item_style_cells[-1]
                    dst_c = sheet2.cell(row=item_row, column=16)
                    if src_c.font: dst_c.font = copy.copy(src_c.font)
                    if src_c.fill: dst_c.fill = copy.copy(src_c.fill)
                    if src_c.alignment: dst_c.alignment = copy.copy(src_c.alignment)
                    
                    # Apply item cell borders
                    sheet2.cell(row=item_row, column=1).border = Border(left=thin_side, right=thin_side, top=thin_side, bottom=thin_side)
                    for col in range(2, 15):
                        sheet2.cell(row=item_row, column=col).border = Border(top=thin_side, bottom=thin_side)
                    sheet2.cell(row=item_row, column=15).border = Border(right=thin_side, top=thin_side, bottom=thin_side)
                    sheet2.cell(row=item_row, column=16).border = Border(left=thin_side, right=thin_side, top=thin_side, bottom=thin_side)
                    
                    # Embed Image
                    img_base64 = item.get('image')
                    has_image = False
                    if img_base64 and ',' in img_base64:
                        try:
                            header, encoded = img_base64.split(',', 1)
                            img_data = base64.b64decode(encoded)
                            
                            ext = ".png"
                            if "jpeg" in header or "jpg" in header:
                                ext = ".jpg"
                                
                            temp_fd, temp_path = tempfile.mkstemp(suffix=ext)
                            os.write(temp_fd, img_data)
                            os.close(temp_fd)
                            temp_files.append(temp_path)
                            
                            with Image.open(temp_path) as pil_img:
                                orig_w, orig_h = pil_img.size
                                
                            max_w, max_h = 380, 180
                            scale = min(max_w / orig_w, max_h / orig_h)
                            new_w = int(orig_w * scale)
                            new_h = int(orig_h * scale)
                            
                            # Center the image inside the cell (Cell P dimensions: 399px width, 200px height)
                            cell_w_px = 399
                            cell_h_px = 200
                            
                            offset_x = max(0, int((cell_w_px - new_w) / 2))
                            offset_y = max(0, int((cell_h_px - new_h) / 2))
                            
                            img = openpyxl.drawing.image.Image(temp_path)
                            img.width = new_w
                            img.height = new_h
                            
                            # Column P is 0-indexed column 15, Row is 0-indexed item_row - 1
                            col_idx = 15
                            row_idx = item_row - 1
                            
                            p2e = pixels_to_EMU
                            col_offset = p2e(offset_x)
                            row_offset = p2e(offset_y)
                            
                            marker = AnchorMarker(col=col_idx, colOff=col_offset, row=row_idx, rowOff=row_offset)
                            size = XDRPositiveSize2D(p2e(new_w), p2e(new_h))
                            
                            img.anchor = OneCellAnchor(_from=marker, ext=size)
                            sheet2.add_image(img)
                            has_image = True
                        except Exception as eImg:
                            print(f"[Excel Generation] Failed to embed image: {eImg}")
                            
                    if has_image:
                        sheet2.row_dimensions[item_row].height = 150.0
                    else:
                        sheet2.row_dimensions[item_row].height = None
                        
                    current_row += 1
                    
            sheet2.print_area = f"A1:P{current_row - 1}"
            
            # Keep ONLY Sheet 2 (Summary Sheet)
            if len(wb.worksheets) > 1:
                wb.remove(wb.worksheets[0])
                
            temp_dir = tempfile.gettempdir()
            out_filename = "ภาพประกอบ.xlsx"
            out_path = os.path.join(temp_dir, out_filename)
            
            wb.save(out_path)
            wb.close()
            
            # Cleanup temp image files
            for f in temp_files:
                try: os.remove(f)
                except: pass
                
            return send_file(out_path, as_attachment=True, download_name=out_filename)

        # ---------------------------------------------
        # Generate Expense Summary Excel (is_illustration = False)
        # ---------------------------------------------
        else:
            # 1. Shift all merged cell ranges starting from row 3 down by 1 row
            ranges = list(sheet2.merged_cells.ranges)
            for r in ranges:
                if r.min_row >= 3:
                    sheet2.merged_cells.remove(r)
                    r.shift(row_shift=1)
                    sheet2.merged_cells.add(r)

            # 2. Insert 1 row at index 3 for the date/location range
            sheet2.insert_rows(3, 1)
            
            # Merge A3:U3 (columns 1 to 21)
            sheet2.merge_cells(start_row=3, start_column=1, end_row=3, end_column=21)
            
            # Copy style from A2 to A3
            sheet2.cell(row=3, column=1).font = copy.copy(sheet2.cell(row=2, column=1).font)
            sheet2.cell(row=3, column=1).alignment = copy.copy(sheet2.cell(row=2, column=1).alignment)
            
            # 3. Cache styles from the shifted template rows (after insertion)
            vendor_style_cells = [sheet2.cell(row=6, column=c) for c in range(1, 27)]
            item_style_cells = [sheet2.cell(row=7, column=c) for c in range(1, 27)]
            discount_style_cells = [sheet2.cell(row=27, column=c) for c in range(1, 27)]
            total_style_cells = [sheet2.cell(row=47, column=c) for c in range(1, 27)]
            
            vendor_row_height = sheet2.row_dimensions[6].height
            item_row_height = sheet2.row_dimensions[7].height
            discount_row_height = sheet2.row_dimensions[27].height
            total_row_height = sheet2.row_dimensions[47].height
            
            # Clear merged cells in dynamic area (starting at row 6 now)
            ranges_to_remove = [r for r in list(sheet2.merged_cells.ranges) if r.min_row >= 6]
            for r in ranges_to_remove:
                sheet2.merged_cells.remove(r)
                
            # Delete template data rows (starting at row 6, deleting everything after)
            sheet2.delete_rows(6, sheet2.max_row - 5)
            
            # Increase Column N width to prevent #####
            sheet2.column_dimensions['N'].width = 15.0
            
            # Update titles
            course_name = data.get('intro_course', '').strip()
            dept = data.get('department', 'สคร.').strip()
            date_range = data.get('excel_date_range', '').strip()
            location = data.get('excel_location', '').strip()
            
            # Row 2: รายการจัดซื้อวัสดุอุปกรณ์สำหรับการจัด หลักสูตร / การดำเนินงาน ...
            if course_name:
                clean_course = course_name.strip(' "\'')
                if clean_course.startswith("รายการจัดซื้อวัสดุอุปกรณ์"):
                    sheet2['A2'] = clean_course
                else:
                    has_prefix = any(clean_course.startswith(prefix) for prefix in [
                        "สำหรับการจัด", "สำหรับ", "สนับสนุน", "เพื่อ", "การจัด"
                    ])
                    if has_prefix:
                        sheet2['A2'] = f"รายการจัดซื้อวัสดุอุปกรณ์{clean_course}"
                    elif clean_course.startswith("ดำเนินงาน"):
                        sheet2['A2'] = f"รายการจัดซื้อวัสดุอุปกรณ์สำหรับการ{clean_course}"
                    else:
                        sheet2['A2'] = f"รายการจัดซื้อวัสดุอุปกรณ์สำหรับการจัด หลักสูตร {clean_course}"
            else:
                sheet2['A2'] = "รายการจัดซื้อวัสดุอุปกรณ์"
                
            # Row 3: วันที่ / ระหว่างวันที่ ... ณ ...
            date_loc_text = ""
            if date_range:
                if date_range.startswith("ระหว่างวันที่") or date_range.startswith("วันที่"):
                    date_loc_text += date_range
                else:
                    has_range_indicator = any(sep in date_range for sep in ["-", "ถึง", "–", "—"])
                    if has_range_indicator:
                        date_loc_text += f"ระหว่างวันที่ {date_range}"
                    else:
                        date_loc_text += f"วันที่ {date_range}"
            if location:
                if date_loc_text:
                    date_loc_text += f" ณ {location}"
                else:
                    date_loc_text += f"ณ {location}"
            sheet2['A3'] = date_loc_text
            
            first_vendor_row = current_row = 6
            
            def apply_cached_style(sheet, row_idx, cached_row_cells):
                for col_idx, src_cell in enumerate(cached_row_cells, 1):
                    dst_cell = sheet.cell(row=row_idx, column=col_idx)
                    if src_cell.font: dst_cell.font = copy.copy(src_cell.font)
                    if src_cell.fill: dst_cell.fill = copy.copy(src_cell.fill)
                    if src_cell.alignment: dst_cell.alignment = copy.copy(src_cell.alignment)
                    if src_cell.border: dst_cell.border = copy.copy(src_cell.border)
                    if src_cell.number_format: dst_cell.number_format = src_cell.number_format
                        
            # Write invoices sorted chronologically by date
            invoices = data.get('invoices', [])
            invoices = sorted(invoices, key=lambda x: parse_thai_date(x.get('invoice_date')))
            for inv_idx, inv in enumerate(invoices, 1):
                items = inv.get('items', [])
                discount = float(inv.get('discount', 0.0))
                
                vendor_header_row = current_row
                start_item_row = current_row + 1
                end_item_row = current_row + len(items)
                
                # Write Vendor Header Row
                sheet2.cell(row=vendor_header_row, column=1, value=inv_idx)
                sheet2.cell(row=vendor_header_row, column=2, value=inv.get('vendor_name', ''))
                sheet2.cell(row=vendor_header_row, column=15, value="=")
                
                # Vendor total formula
                if discount > 0:
                    discount_row_idx = end_item_row + 1
                    vendor_total_formula = f"=SUM(N{start_item_row}:N{end_item_row})-E{discount_row_idx}"
                else:
                    vendor_total_formula = f"=SUM(N{start_item_row}:N{end_item_row})"
                    
                sheet2.cell(row=vendor_header_row, column=16, value=vendor_total_formula)
                sheet2.cell(row=vendor_header_row, column=17, value="บาท")
                sheet2.cell(row=vendor_header_row, column=18, value=f"=P{vendor_header_row}")
                
                apply_cached_style(sheet2, vendor_header_row, vendor_style_cells)
                sheet2.row_dimensions[vendor_header_row].height = vendor_row_height
                
                # Recreate merges
                sheet2.merge_cells(start_row=vendor_header_row, start_column=2, end_row=vendor_header_row, end_column=14)
                sheet2.merge_cells(start_row=vendor_header_row, start_column=18, end_row=vendor_header_row, end_column=20)
                
                current_row += 1
                
                # Write Item Rows
                for item_idx, item in enumerate(items, 1):
                    item_row = current_row
                    
                    code = item.get('item_code', '').strip()
                    desc = item.get('description', '').strip()
                    desc_val = f"{code} {desc}" if code else desc
                    
                    sheet2.cell(row=item_row, column=1, value=f"{inv_idx}.{item_idx}")
                    sheet2.cell(row=item_row, column=2, value=desc_val)
                    sheet2.cell(row=item_row, column=4, value="(")
                    sheet2.cell(row=item_row, column=5, value=float(item.get('unit_price', 0.0)))
                    sheet2.cell(row=item_row, column=6, value="บาท * ")
                    sheet2.cell(row=item_row, column=7, value=int(item.get('quantity', 1)))
                    sheet2.cell(row=item_row, column=8, value=item.get('unit', 'ชิ้น'))
                    sheet2.cell(row=item_row, column=9, value=")")
                    sheet2.cell(row=item_row, column=13, value="=")
                    sheet2.cell(row=item_row, column=14, value=f"=E{item_row}*G{item_row}")
                    sheet2.cell(row=item_row, column=15, value="บาท")
                    
                    apply_cached_style(sheet2, item_row, item_style_cells)
                    from openpyxl.styles import Alignment
                    desc_c = sheet2.cell(row=item_row, column=2)
                    desc_c.alignment = Alignment(wrapText=True, vertical='top', horizontal='left')
                    sheet2.merge_cells(start_row=item_row, start_column=2, end_row=item_row, end_column=3)
                    sheet2.row_dimensions[item_row].height = None
                    
                    current_row += 1
                    
                # Write Discount Row
                if discount > 0:
                    discount_row = current_row
                    sheet2.cell(row=discount_row, column=2, value="ส่วนลด")
                    sheet2.cell(row=discount_row, column=5, value=discount)
                    sheet2.cell(row=discount_row, column=6, value="บาท")
                    
                    apply_cached_style(sheet2, discount_row, discount_style_cells)
                    sheet2.row_dimensions[discount_row].height = discount_row_height
                    
                    sheet2.merge_cells(start_row=discount_row, start_column=2, end_row=discount_row, end_column=3)
                    current_row += 1
                    
                # Merge Column U vertically
                merge_start_row = start_item_row
                merge_end_row = current_row - 1
                if merge_end_row > merge_start_row:
                    sheet2.merge_cells(start_row=merge_start_row, start_column=21, end_row=merge_end_row, end_column=21)
                    
            # Write Total Row
            total_row = current_row
            sheet2.cell(row=total_row, column=1, value="รวม")
            sheet2.cell(row=total_row, column=2, value=f"=BAHTTEXT(R{total_row})")
            sheet2.cell(row=total_row, column=18, value=f"=SUM(R{first_vendor_row}:R{total_row - 1})")
            
            apply_cached_style(sheet2, total_row, total_style_cells)
            sheet2.row_dimensions[total_row].height = total_row_height
            
            sheet2.merge_cells(start_row=total_row, start_column=2, end_row=total_row, end_column=17)
            sheet2.merge_cells(start_row=total_row, start_column=18, end_row=total_row, end_column=20)
            
            # Ensure outer left/right borders are medium
            from openpyxl.styles import Border, Side
            medium_side = Side(border_style="medium", color="000000")
            for r in range(5, total_row + 1):
                c_a = sheet2.cell(row=r, column=1)
                c_a.border = Border(
                    left=medium_side,
                    right=c_a.border.right if c_a.border else None,
                    top=c_a.border.top if c_a.border else None,
                    bottom=c_a.border.bottom if c_a.border else None
                )
                c_u = sheet2.cell(row=r, column=21)
                c_u.border = Border(
                    left=c_u.border.left if c_u.border else None,
                    right=medium_side,
                    top=c_u.border.top if c_u.border else None,
                    bottom=c_u.border.bottom if c_u.border else None
                )
                
            sheet2.print_area = f"A1:U{total_row}"
            
            # Page setup and print options for summary sheet
            sheet2.page_setup.orientation = sheet2.ORIENTATION_PORTRAIT
            sheet2.page_setup.paperSize = sheet2.PAPERSIZE_A4
            sheet2.page_setup.fitToWidth = 1
            sheet2.page_setup.fitToHeight = 0
            sheet2.sheet_properties.pageSetUpPr.fitToPage = True
            sheet2.print_options.gridLines = True
            if sheet2.views.sheetView:
                for view in sheet2.views.sheetView:
                    view.showGridLines = True
            else:
                sheet2.views.sheetView.append(openpyxl.worksheet.views.SheetView(showGridLines=True))

            # Keep ONLY Sheet 2 (Summary Sheet)
            if len(wb.worksheets) > 1:
                wb.remove(wb.worksheets[0])
                
            temp_dir = tempfile.gettempdir()
            out_filename = "สรุปค่าใช้จ่าย_เบิกเงินค่าพัสดุ.xlsx"
            out_path = os.path.join(temp_dir, out_filename)
            
            wb.save(out_path)
            wb.close()
            
            return send_file(out_path, as_attachment=True, download_name=out_filename)
            
    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({"error": f"Failed during Excel generation: {str(e)}"}), 500

if __name__ == "__main__":
    # Run server on port 5000, allowing local network access (host="0.0.0.0")
    app.run(host="0.0.0.0", port=5000, debug=True)

