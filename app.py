import os
import sys
import tempfile
import copy
import re
import time
from flask import Flask, request, jsonify, send_file
from PIL import Image
from google import genai
from google.genai import types
from pydantic import BaseModel, Field
from typing import List, Optional
import docx
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from bahttext import bahttext


app = Flask(__name__, static_folder='static', static_url_path='')

# Define Schema for Gemini Bill Extraction
class BillItem(BaseModel):
    item_code: Optional[str] = Field(None, description="The product code, SKU, or barcode (e.g. MS4-000939). If not available, leave null.")
    description: str = Field(description="Detailed item name or description in Thai or English as written on the bill. Keep serial numbers (S/N) if present.")
    quantity: int = Field(description="Quantity purchased")
    unit_price: float = Field(description="Unit price of the item")
    total_price: float = Field(description="Total price of the item (quantity * unit_price)")

class BillExtraction(BaseModel):
    vendor_name: str = Field(description="Name of the company or store that issued the bill/receipt (in Thai, e.g. 'บริษัท ไอที ซิตี้ (มหาชน)' or 'บริษัท บีทูเอส จำกัด')")
    invoice_number: str = Field(description="Invoice or Receipt number (เลขที่/เลขที่ใบเสร็จ/เล่มที่-เลขที่) as it appears on the document")
    invoice_date: str = Field(description="Date of the invoice/receipt in Thai format (e.g. '28 ตุลาคม 2567')")
    items: List[BillItem] = Field(description="List of all purchased items")
    discount: float = Field(0.0, description="Total discount applied to the bill if any")
    grand_total: float = Field(description="Grand total / total amount paid in Baht")

# Font and formatting helpers for TH SarabunPSK
def set_run_font_enhanced(run, font_name='TH SarabunPSK', font_size_pt=16, bold=False, italic=False, underline=None, spacing=None):
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(font_size_pt)
    run.font.name = font_name
    
    if underline is not None:
        run.font.underline = underline
        
    rPr = run._r.get_or_add_rPr()
    
    # Set explicit font size in XML (sz and szCs in half-points)
    sz = rPr.find(qn('w:sz'))
    if sz is None:
        sz = OxmlElement('w:sz')
        rPr.append(sz)
    sz.set(qn('w:val'), str(int(font_size_pt * 2)))
    
    szCs = rPr.find(qn('w:szCs'))
    if szCs is None:
        szCs = OxmlElement('w:szCs')
        rPr.append(szCs)
    szCs.set(qn('w:val'), str(int(font_size_pt * 2)))
    
    if bold:
        # Set bold complex script
        bCs = rPr.find(qn('w:bCs'))
        if bCs is None:
            bCs = OxmlElement('w:bCs')
            rPr.append(bCs)
    if italic:
        # Set italic complex script
        iCs = rPr.find(qn('w:iCs'))
        if iCs is None:
            iCs = OxmlElement('w:iCs')
            rPr.append(iCs)
        
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rFonts.set(qn('w:cs'), font_name)
    rFonts.set(qn('w:eastAsia'), font_name)
    
    # Set character spacing (w:spacing) in dxa (dxa = 1/20 pt)
    if spacing is not None:
        spacing_elem = rPr.find(qn('w:spacing'))
        if spacing_elem is None:
            spacing_elem = OxmlElement('w:spacing')
            rPr.append(spacing_elem)
        spacing_elem.set(qn('w:val'), str(spacing))

def set_run_font(run, font_name='TH SarabunPSK', font_size_pt=16, bold=False, italic=False):
    set_run_font_enhanced(run, font_name, font_size_pt, bold, italic)

def copy_paragraph_format(src_p, dst_p):
    # Get the paragraph properties XML elements
    src_pPr = src_p._p.get_or_add_pPr()
    dst_pPr = dst_p._p.get_or_add_pPr()
    
    # Clear existing paragraph properties
    dst_pPr.clear()
    
    # Deep copy all child elements of paragraph properties (alignment, spacing, indents, styles, tab stops)
    for child in src_pPr:
        dst_pPr.append(copy.deepcopy(child))

def format_price(val):
    if val % 1 == 0:
        return f"{val:,.0f}"
    else:
        return f"{val:,.2f}"

def update_paragraph_3(p, dept, phone):
    # Change department run
    for r in p.runs:
        if "สคร." in r.text:
            r.text = dept
            break
            
    # Change phone run
    phone_run_idx = None
    for idx, r in enumerate(p.runs):
        if "033 005 833" in r.text:
            phone_run_idx = idx
            break
            
    if phone_run_idx is not None:
        p.runs[phone_run_idx].text = phone
        # Clear the old "(ณัฐภัท)" runs instead of deleting them to preserve trailing tabs
        for i in range(phone_run_idx + 1, len(p.runs)):
            text_to_check = p.runs[i].text
            if any(char in text_to_check for char in ['(', 'ณัฐ', 'ภัท', ')']):
                p.runs[i].text = ""

def update_paragraph_4(p, memo_no, date_val):
    # Find the run containing "สคร."
    memo_start_idx = None
    for idx, r in enumerate(p.runs):
        if "สคร." in r.text:
            memo_start_idx = idx
            break
            
    # Find the run containing "วันที่"
    date_label_idx = None
    for idx, r in enumerate(p.runs):
        if "วันที่" in r.text:
            date_label_idx = idx
            break
            
    if memo_start_idx is not None and date_label_idx is not None:
        # Put the entire memo_no in the memo_start_idx run
        p.runs[memo_start_idx].text = memo_no
        # Clear the runs between memo_start_idx + 1 and date_label_idx - 1 (preserves spaces before "วันที่")
        for i in range(memo_start_idx + 1, date_label_idx - 1):
            p.runs[i].text = ""
            
    # Find the run containing "พฤศจิกายน"
    date_val_idx = None
    for idx, r in enumerate(p.runs):
        if "พฤศจิกายน" in r.text:
            date_val_idx = idx
            break
            
    if date_val_idx is not None:
        # Put the date value in the date_val_idx run
        p.runs[date_val_idx].text = f" {date_val} "
        # Clear the run containing "  2567" (which is typically date_val_idx + 1)
        if date_val_idx + 1 < len(p.runs) and "2567" in p.runs[date_val_idx + 1].text:
            p.runs[date_val_idx + 1].text = ""

def update_paragraph_5(p, subject):
    # Find "เรื่อง"
    subject_label_idx = None
    for idx, r in enumerate(p.runs):
        if "เรื่อง" in r.text:
            subject_label_idx = idx
            break
            
    if subject_label_idx is not None and subject_label_idx + 2 < len(p.runs):
        # Put the new subject text in Run 2 (subject_label_idx + 2)
        p.runs[subject_label_idx + 2].text = subject
        # Clear Runs 3 to 8 (which contained old subject portions)
        for i in range(subject_label_idx + 3, subject_label_idx + 9):
            if i < len(p.runs):
                p.runs[i].text = ""

def update_paragraph_6(p, to_text):
    # Find "เรียน"
    learn_idx = None
    for idx, r in enumerate(p.runs):
        if "เรียน" in r.text:
            learn_idx = idx
            break
            
    if learn_idx is not None and learn_idx + 2 < len(p.runs):
        # Put the new to_text in Run 2 (learn_idx + 2)
        p.runs[learn_idx + 2].text = to_text
        # Clear Runs 3 to 6
        for i in range(learn_idx + 3, len(p.runs)):
            p.runs[i].text = ""

def render_intro_paragraph(p, intro_text, dept, total_items):
    # Try parsing intro_text using regex to extract variables
    pattern = r"ด้วย\s+(.*?)\s+ได้ดำเนินการจัดซื้อวัสดุสำหรับการจัด\s+หลักสูตร\s+(.*?)\s+จำนวน\s+(.*?)\s+รายการ\s+โดยมีวัตถุประสงค์ตาม\s+(.*?)\s+โดยใช้งบประมาณ\s+\(รหัสงบประมาณ\s+(.*?)\s*ซึ่งมีรายละเอียดดังต่อไปนี้"
    match = re.search(pattern, intro_text, re.DOTALL)
    
    if match and len(p.runs) >= 41:
        val_dept = match.group(1).strip()
        val_course = match.group(2).strip()
        val_qty = match.group(3).strip()
        val_project = match.group(4).strip()
        val_budget = match.group(5).strip()
        
        # 1. Dept (Run 2)
        p.runs[2].text = val_dept
        
        # 2. Course (Run 4 and 5)
        p.runs[4].text = "ได้ดำเนินการจัดซื้อวัสดุสำหรับการจัด หลักสูตร " + val_course
        p.runs[5].text = ""
        
        # 3. Qty (Run 7)
        p.runs[7].text = str(total_items)
        
        # 4. Project (Run 12 to 20)
        p.runs[12].text = val_project
        for i in range(13, 21):
            p.runs[i].text = ""
            
        # 5. Budget (Run 24 to 38)
        p.runs[24].text = " " + val_budget
        for i in range(25, 39):
            p.runs[i].text = ""
        return
        
    # Fallback
    p.text = ""
    run = p.add_run(intro_text)
    set_run_font_enhanced(run, 'TH SarabunPSK', 16)

def update_regulatory_paragraph(p, reg_text):
    # Search for list index value at the end
    match = re.search(r"ตาราง\s+1\s+ลำดับที่\s+(.*)", reg_text)
    if match and len(p.runs) >= 13:
        idx_val = match.group(1).strip()
        # Set Run 12 to index value (which has underline)
        p.runs[12].text = f" {idx_val}"
        # Set Run 11 to empty space
        p.runs[11].text = ""
        return
        
    # Fallback: recreate paragraph runs manually
    p.text = ""
    run = p.add_run(reg_text)
    set_run_font_enhanced(run, 'TH SarabunPSK', 16)

def add_item_paragraph(doc, current_anchor, ref_item_p, idx, item):
    new_p = doc.add_paragraph()
    current_anchor._element.addnext(new_p._element)
    copy_paragraph_format(ref_item_p, new_p)
    
    desc_val = item["desc"]
    if item["code"]:
        desc_val = f"{item['code']} {desc_val}"
        
    price_str = format_price(item["price"])
    
    # Precise run structures and spacings to match paragraph 8 of the template exactly
    runs_data = [
        ("\t", False, -6),
        (f"{idx}. ค่า", False, -6),
        ("\t", True, -6),
        (desc_val, True, -6),
        (" ", True, -6),
        ("จำนวน", False, -6),
        ("  ", True, -6),
        (str(item["qty"]), True, -6),
        ("  ", True, -6),
        ("รายการ เป็นเงิน", False, -6),
        ("  ", True, -6),
        (price_str, True, -6),
        (" ", True, -6),
        (" ", True, -6),
        ("บาท", False, -6),
        ("จากบริษัท", False, -6),
        (" ", False, -6),
        ("\t", True, None),
        (f"{item['vendor']}", True, -6),
        ("   ", True, None),
        ("ตามหลักฐานการจัดซื้อเป็น", False, -6),
        (f"  {item['doc_type']}", True, -6),
        (" ", False, -6),
        ("เล่มที่-เลขที่ ", False, -6),
        (" ", True, -6),
        (item["invoice_no"], True, -6),
        ("  ", True, -6),
        ("วันที่", False, -6),
        (" ", True, -6),
        (f"{item['invoice_date']}", True, -6),
        ("  ", True, -6)
    ]
    
    for text_val, is_underlined, spacing_val in runs_data:
        run = new_p.add_run(text_val)
        underline_style = docx.enum.text.WD_UNDERLINE.DOTTED if is_underlined else None
        set_run_font_enhanced(run, 'TH SarabunPSK', 16, underline=underline_style, spacing=spacing_val)
        
    return new_p

def add_note_paragraph(doc, current_anchor, ref_note_p, start_num, end_num, subtotal, discount):
    new_p = doc.add_paragraph()
    current_anchor._element.addnext(new_p._element)
    copy_paragraph_format(ref_note_p, new_p)
    
    subtotal_str = format_price(subtotal)
    discount_str = format_price(discount)
    
    note_text = f"หมายเหตุ : รายการที่ {start_num} – {end_num} ยอดรวม {subtotal_str} บาท ได้รับส่วนลดทั้งหมด เป็นเงิน {discount_str} บาท    "
    
    run = new_p.add_run(note_text)
    set_run_font_enhanced(run, 'TH SarabunPSK', 16, underline=docx.enum.text.WD_UNDERLINE.DOTTED, spacing=-4)
    
    return new_p

def update_total_paragraph(total_p, total_items, grand_total, thai_text):
    # Clear runs
    for r in list(total_p.runs):
        total_p._p.remove(r._r)
        
    grand_total_str = format_price(grand_total)
    
    runs_data = [
        ("\t", False, None),
        ("รวม", False, None),
        ("\t", True, None),
        (str(total_items), True, None),
        ("\t", True, None),
        ("รายการ เป็นเงินทั้งสิ้น", False, None),
        ("\t", True, None),
        (grand_total_str, True, None),
        ("  บาท", True, None),
        ("\t", True, None),
        (f"({thai_text})", True, None),
        ("\t", True, None),
        ("  ", True, None)
    ]
    
    for text_val, is_underlined, spacing_val in runs_data:
        run = total_p.add_run(text_val)
        underline_style = docx.enum.text.WD_UNDERLINE.DOTTED if is_underlined else None
        set_run_font_enhanced(run, 'TH SarabunPSK', 16, underline=underline_style, spacing=spacing_val)

@app.route('/')
def index():
    return app.send_static_file('index.html')

@app.route('/api/validate_key', methods=['POST'])
def validate_key():
    data = request.json or {}
    api_key = data.get('api_key') or request.headers.get('Authorization')
    if api_key and api_key.startswith('Bearer '):
        api_key = api_key[len('Bearer '):]
    
    if not api_key:
        return jsonify({"valid": False, "error": "กรุณากรอก API Key"}), 400
        
    try:
        client = genai.Client(api_key=api_key)
        
        # Test content generation with the primary models to ensure the key actually works
        try:
            client.models.generate_content(
                model='gemini-2.5-flash',
                contents='Hello'
            )
        except Exception as e25:
            # Fallback test with gemini-1.5-flash
            client.models.generate_content(
                model='gemini-1.5-flash',
                contents='Hello'
            )
            
        return jsonify({
            "valid": True, 
            "message": "API Key สามารถใช้งานได้ปกติ!"
        })
    except Exception as e:
        err_msg = str(e)
        import traceback
        traceback.print_exc()
        
        friendly_err = err_msg
        tip = "กรุณาสร้าง API Key ใหม่จาก Google AI Studio (https://aistudio.google.com/)"
        
        if "PERMISSION_DENIED" in err_msg or "denied" in err_msg.lower():
            friendly_err = "สิทธิ์การเข้าใช้งานถูกปฏิเสธ (403 Permission Denied) คีย์หรือโปรเจกต์ของคุณถูกระงับสิทธิ์เข้าใช้งานระบบจาก Google"
        elif "API_KEY_INVALID" in err_msg or "key is invalid" in err_msg.lower():
            friendly_err = "API Key ไม่ถูกต้อง (400 Invalid API Key) กรุณาตรวจสอบและคัดลอกคีย์ใหม่อีกครั้ง"
        elif "NOT_FOUND" in err_msg or "not found" in err_msg.lower():
            friendly_err = "ไม่พบบริการโมเดล (404 Not Found) บริการ Google Generative Language API ยังไม่เปิดใช้งานในโปรเจกต์ของคีย์นี้"
            tip = "หากใช้ GCP API Key กรุณาเข้าไปเปิดใช้งาน 'Generative Language API' ใน Google Cloud Console หรือแนะนำให้สร้างคีย์ฟรีจาก Google AI Studio แทน"
            
        return jsonify({
            "valid": False,
            "error": friendly_err,
            "tip": tip
        })


def call_gemini_with_retries(client, img, prompt, schema=None):
    models_to_try = ['gemini-2.5-flash', 'gemini-2.0-flash', 'gemini-1.5-flash']
    
    # Try with schema first if requested
    if schema:
        for model_name in models_to_try:
            for attempt in range(3):
                try:
                    print(f"Attempting extraction with {model_name} (with schema), attempt {attempt+1}...")
                    response = client.models.generate_content(
                        model=model_name,
                        contents=[img, prompt],
                        config=types.GenerateContentConfig(
                            response_mime_type="application/json",
                            response_schema=schema,
                        ),
                    )
                    return response.text
                except Exception as e:
                    err_msg = str(e)
                    print(f"{model_name} (with schema) attempt {attempt+1} failed: {err_msg}")
                    # If it's a 503 or 429, wait and retry
                    if any(code in err_msg for code in ["503", "429"]) or any(kw in err_msg.lower() for kw in ["temporarily", "limit", "demand", "busy", "unavailable"]):
                        time.sleep(2 ** attempt)  # 1s, 2s, 4s backoff
                    else:
                        break # Go to next model immediately for 404, 403, etc.
                        
    # Fallback to plain text JSON prompts
    plain_prompt = prompt + "\nRespond ONLY with a valid JSON matching this schema: {\"vendor_name\": \"...\", \"invoice_number\": \"...\", \"invoice_date\": \"...\", \"items\": [{\"item_code\": \"...\", \"description\": \"...\", \"quantity\": 1, \"unit_price\": 1.0, \"total_price\": 1.0}], \"discount\": 0.0, \"grand_total\": 0.0}. Do not include markdown formatting or backticks."
    
    for model_name in models_to_try:
        for attempt in range(3):
            try:
                print(f"Attempting extraction with {model_name} (plain JSON), attempt {attempt+1}...")
                response = client.models.generate_content(
                    model=model_name,
                    contents=[img, plain_prompt]
                )
                text = response.text.strip()
                # Clean up markdown code block if present
                if text.startswith("```"):
                    lines = text.split("\n")
                    if lines[0].startswith("```"):
                        lines = lines[1:]
                    if lines[-1].startswith("```"):
                        lines = lines[:-1]
                    text = "\n".join(lines).strip()
                return text
            except Exception as e:
                err_msg = str(e)
                print(f"{model_name} (plain JSON) attempt {attempt+1} failed: {err_msg}")
                if any(code in err_msg for code in ["503", "429"]) or any(kw in err_msg.lower() for kw in ["temporarily", "limit", "demand", "busy", "unavailable"]):
                    time.sleep(2 ** attempt)
                else:
                    break
                    
    raise Exception("All models and formats failed. Please check your API key / Google Cloud settings, or try again later.")

@app.route('/api/extract', methods=['POST'])
def extract_bill():
    api_key = request.headers.get('Authorization')
    if api_key and api_key.startswith('Bearer '):
        api_key = api_key[len('Bearer '):]
    
    if not api_key:
        api_key = request.form.get('api_key')
        
    if not api_key:
        return jsonify({"error": "Gemini API key is required"}), 400
        
    if 'file' not in request.files:
        return jsonify({"error": "No file uploaded"}), 400
        
    file = request.files['file']
    if file.filename == '':
        return jsonify({"error": "No file selected"}), 400
        
    try:
        img = Image.open(file.stream)
        client = genai.Client(api_key=api_key)
        
        prompt = """
        Extract the bill information in detail. Make sure to:
        1. Identify the vendor name, invoice/receipt number, and invoice date (convert to Thai format e.g. "28 ตุลาคม 2567").
        2. List every item in the bill with its code (SKU/barcode if available), description, quantity, unit price, and total price.
        3. If there is a product code (e.g. MS4-000939 or a barcode V 885...), place it in item_code.
        4. Extract any discount and the grand total.
        """
        
        response_text = call_gemini_with_retries(client, img, prompt, schema=BillExtraction)
        return response_text, 200, {'Content-Type': 'application/json'}
        
    except Exception as e:
        import traceback
        traceback.print_exc()
        
        err_msg = str(e)
        tip_msg = ""
        if "NOT_FOUND" in err_msg or "not found" in err_msg.lower() or "404" in err_msg:
            tip_msg = "\n\n💡 Tip: This usually means the Generative Language API is not enabled for your project, or your API key is invalid. Please create a key from Google AI Studio (https://aistudio.google.com/) and try again."
            
        return jsonify({"error": f"Failed during bill extraction: {err_msg}{tip_msg}"}), 500

@app.route('/api/generate', methods=['POST'])
def generate_docx():
    data = request.json
    if not data:
        return jsonify({"error": "Missing payload"}), 400
        
    template_path = r"d:\AutoWord\แบบฟอร์ม รายงานขอความเห็นชอบซื้อจ้าง.docx"
    if not os.path.exists(template_path):
        return jsonify({"error": "Template document file not found on server"}), 500
        
    try:
        doc = docx.Document(template_path)
        
        # Process items to list first to get total_items_count
        invoices = data.get('invoices', [])
        flat_items = []
        item_counter = 1
        invoice_ranges = []
        
        for inv in invoices:
            items_list = inv.get('items', [])
            if not items_list:
                continue
            start_num = item_counter
            for item in items_list:
                flat_items.append({
                    "code": item.get('item_code'),
                    "desc": item.get('description'),
                    "qty": item.get('quantity', 1),
                    "price": item.get('total_price', 0.0),
                    "vendor": inv.get('vendor_name'),
                    "invoice_no": inv.get('invoice_number'),
                    "invoice_date": inv.get('invoice_date'),
                    "doc_type": inv.get('doc_type', 'ใบเสร็จรับเงิน/ใบกำกับภาษี')
                })
                item_counter += 1
            end_num = item_counter - 1
            
            invoice_ranges.append({
                "start": start_num,
                "end": end_num,
                "discount": inv.get('discount', 0.0),
                "subtotal": sum(item.get('total_price', 0.0) for item in items_list),
                "vendor": inv.get('vendor_name')
            })
            
        total_items_count = len(flat_items)
        
        # 1. Update Header Paragraphs
        # Paragraph 3: Department & Phone
        update_paragraph_3(doc.paragraphs[3], data.get('department', 'สคร.'), data.get('phone', ''))
        
        # Paragraph 4: Memo No & Date
        update_paragraph_4(doc.paragraphs[4], data.get('memo_no', 'สคร.             /2567'), data.get('date', 'พฤศจิกายน  2567'))
        
        # Paragraph 5: Subject
        subject_text = f"รายงานขอความเห็นชอบการจัดซื้อจัดจ้าง  จำนวน {total_items_count}  รายการ"
        update_paragraph_5(doc.paragraphs[5], data.get('subject', subject_text))
        
        # Paragraph 6: Learn (เรียน)
        update_paragraph_6(doc.paragraphs[6], data.get('to_text', 'ผอ.สคร.  ผ่าน รก.หน.ฝถท.'))
        
        # Paragraph 7: Introduction Body
        render_intro_paragraph(doc.paragraphs[7], data.get('intro_text', ''), data.get('department', 'สคร.'), total_items_count)
        
        # 2. Find and Replace the Items List
        start_idx = None
        end_idx = None
        
        for i, p in enumerate(doc.paragraphs):
            text = p.text.strip()
            if text.startswith("1. ค่า"):
                start_idx = i
                break
                
        for i, p in enumerate(doc.paragraphs):
            text = p.text.strip()
            if text.startswith("รวม") and "รายการ เป็นเงินทั้งสิ้น" in text:
                end_idx = i
                break
                
        if start_idx is None or end_idx is None:
            return jsonify({"error": "Invalid template format: Could not find items section"}), 500
            
        ref_item_p = doc.paragraphs[start_idx]
        ref_note_p = doc.paragraphs[start_idx + 1] if "หมายเหตุ" in doc.paragraphs[start_idx + 1].text else ref_item_p
        
        # Delete old items list
        paragraphs_to_delete = doc.paragraphs[start_idx:end_idx]
        for p in paragraphs_to_delete:
            p_elem = p._element
            p_elem.getparent().remove(p_elem)
            p._element = None
            p._p = None
            
        # Re-fetch anchor paragraph
        anchor_p = doc.paragraphs[start_idx - 1]
        
        # Generate new paragraphs
        current_anchor = anchor_p
        
        # Track which invoice range a flat item belongs to so we can insert the correct note
        for idx, item in enumerate(flat_items, 1):
            # 1. Add Item Paragraph
            new_item_p = add_item_paragraph(doc, current_anchor, ref_item_p, idx, item)
            current_anchor = new_item_p
            
            # 2. Find if this item belongs to an invoice that has a discount
            matching_range = None
            for r in invoice_ranges:
                if r["start"] <= idx <= r["end"]:
                    matching_range = r
                    break
                    
            if matching_range and matching_range["discount"] > 0:
                # Add note paragraph
                new_note_p = add_note_paragraph(doc, current_anchor, ref_note_p, matching_range['start'], matching_range['end'], matching_range['subtotal'], matching_range['discount'])
                current_anchor = new_note_p
                
        # 3. Update Total Paragraph (It moved down, let's find the new index)
        total_p_idx = None
        for i, p in enumerate(doc.paragraphs):
            if p.text.strip().startswith("รวม") and "รายการ เป็นเงินทั้งสิ้น" in p.text:
                total_p_idx = i
                break
                
        if total_p_idx is not None:
            total_p = doc.paragraphs[total_p_idx]
            
            # Calculate grand total sums
            grand_subtotal = sum(item["price"] for item in flat_items)
            total_discount = sum(r["discount"] for r in invoice_ranges)
            grand_total_paid = grand_subtotal - total_discount
            
            thai_text_amount = bahttext(grand_total_paid)
            
            # Update total paragraph
            update_total_paragraph(total_p, total_items_count, grand_total_paid, thai_text_amount)
            
        # 4. Update regulatory note
        reg_p_idx = total_p_idx + 1 if total_p_idx is not None else None
        if reg_p_idx is not None and reg_p_idx < len(doc.paragraphs):
            reg_p = doc.paragraphs[reg_p_idx]
            update_regulatory_paragraph(reg_p, data.get('regulatory_text', ''))
            
        # 5. Update signature blocks
        # Replace names and positions in the remaining paragraphs
        for p in doc.paragraphs[reg_p_idx + 1:]:
            p_text = p.text
            
            # 5.1 Requester Position (Paragraph 61)
            if "เจ้าหน้าที่/ผู้รับผิดชอบ" in p_text:
                for r in p.runs:
                    if "เจ้าหน้าที่/ผู้รับผิดชอบ" in r.text:
                        r.text = data.get('requester_position', 'เจ้าหน้าที่/ผู้รับผิดชอบ')
                        break
                        
            # 5.2 Requester Name and Date (Paragraph 62)
            elif "ศิริพักตร์" in p_text or "เสมียนคิด" in p_text:
                # Replace name
                for r in p.runs:
                    if "ศิริพักตร์" in r.text or "เสมียนคิด" in r.text:
                        r.text = data.get('requester_name', 'นางสาวศิริพักตร์  เสมียนคิด')
                        break
                # Replace date
                date_start_idx = None
                for idx, r in enumerate(p.runs):
                    if "/" in r.text or "พฤศจิกายน" in r.text:
                        date_start_idx = idx
                        break
                if date_start_idx is not None:
                    # Put new date in the first date run (Run 18)
                    p.runs[date_start_idx].text = data.get('requester_date', '   / พฤศจิกายน / 2567')
                    # Clear subsequent runs (Runs 19 to 23) in-place to preserve trailing tabs
                    for i in range(date_start_idx + 1, date_start_idx + 6):
                        if i < len(p.runs):
                            p.runs[i].text = ""
                    
            # 5.3 Approver Position (Paragraph 66)
            elif "ผอ.สคร." in p_text:
                for r in p.runs:
                    if "ผอ.สคร." in r.text:
                        r.text = data.get('approver_position', 'ผอ.สคร.')
                        break
                        
            # 5.4 Approver Name (Paragraph 67)
            elif "ปราณปริยา" in p_text or "วงค์ษา" in p_text:
                for r in p.runs:
                    if "ปราณปริยา" in r.text or "วงค์ษา" in r.text:
                        r.text = data.get('approver_name', 'นางสาวปราณปริยา   วงค์ษา')
                        break
                        
            # 5.5 Approver Date (Paragraph 68)
            elif ("/ พฤศจิกายน" in p_text or "/" in p_text) and not ("ศิริพักตร์" in p_text or "เสมียนคิด" in p_text):
                date_start_idx = None
                for idx, r in enumerate(p.runs):
                    if "/" in r.text or "พฤศจิกายน" in r.text:
                        date_start_idx = idx
                        break
                if date_start_idx is not None:
                    # Put new date in the first date run (Run 8)
                    p.runs[date_start_idx].text = data.get('approver_date', '   / พฤศจิกายน / 2567')
                    # Clear subsequent runs (Runs 9 and 10) in-place to preserve trailing tabs
                    for i in range(date_start_idx + 1, date_start_idx + 3):
                        if i < len(p.runs):
                            p.runs[i].text = ""
                    
        # Save to temporary file and return
        temp_dir = tempfile.gettempdir()
        out_filename = "รายงานขอความเห็นชอบการจัดซื้อจัดจ้าง.docx"
        out_path = os.path.join(temp_dir, out_filename)
        
        doc.save(out_path)
        
        return send_file(out_path, as_attachment=True, download_name=out_filename)
        
    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({"error": f"Failed during document generation: {str(e)}"}), 500

if __name__ == "__main__":
    # Run server on port 5000
    app.run(host="127.0.0.1", port=5000, debug=True)
